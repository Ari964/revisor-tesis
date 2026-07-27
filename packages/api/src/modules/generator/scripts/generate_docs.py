import sys
import os
import json
import base64
import tempfile
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Flowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, inch
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# ---------------------------------------------------------------------------
# Helper functions for Roman Numerals and XML manipulation
# ---------------------------------------------------------------------------

def to_roman(num):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syb[i]
            num -= val[i]
        i += 1
    return roman_num.lower()

def configure_page_numbering(section, fmt="decimal", start=None):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    if fmt:
        pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))

def get_clean_text(val):
    if val is None:
        return ""
    if isinstance(val, dict):
        text = ""
        if "general" in val:
            text += f"General:\n{val['general']}\n\n"
        especificos_key = None
        for k in ["especificos", "especificas"]:
            if k in val:
                especificos_key = k
                break
        if especificos_key and isinstance(val[especificos_key], list):
            text += "Específicos:\n"
            for item in val[especificos_key]:
                text += f"- {item}\n"
        if not text:
            for k, v in val.items():
                text += f"{k.capitalize()}: {v}\n"
        return text.strip()
    return str(val)

def get_clean_html_text(val):
    if val is None:
        return ""
    if isinstance(val, dict):
        text = ""
        if "general" in val:
            text += f"<b>General:</b><br/>{val['general']}<br/><br/>"
        especificos_key = None
        for k in ["especificos", "especificas"]:
            if k in val:
                especificos_key = k
                break
        if especificos_key and isinstance(val[especificos_key], list):
            text += "<b>Específicos:</b><br/>"
            for item in val[especificos_key]:
                text += f"• {item}<br/>"
        if not text:
            for k, v in val.items():
                text += f"<b>{k.capitalize()}:</b> {v}<br/>"
        return text
    return str(val)

def add_page_number_to_footer(paragraph):
    p = paragraph._p
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    p.append(fldSimple)

def parse_markdown_to_elements(text):
    if not text:
        return []
    lines = [line.strip() for line in text.split('\n')]
    elements = []
    in_table = False
    table_lines = []
    
    for line in lines:
        if line.startswith('|') and line.endswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                table_element = parse_markdown_table(table_lines)
                if table_element:
                    elements.append(table_element)
                table_lines = []
                in_table = False
                
        if not line:
            continue
            
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            txt = line.lstrip('#').strip()
            elements.append(('h', level, txt))
        elif line.startswith(('*', '-', '+')) and len(line) > 1:
            txt = line[1:].strip()
            elements.append(('li', txt))
        elif line.startswith(tuple(f"{i}." for i in range(1, 100))):
            parts = line.split('.', 1)
            txt = parts[1].strip()
            elements.append(('li', txt))
        else:
            elements.append(('p', line))
            
    if in_table and table_lines:
        table_element = parse_markdown_table(table_lines)
        if table_element:
            elements.append(table_element)
            
    return elements

def parse_markdown_table(lines):
    if len(lines) < 1:
        return None
    parsed_rows = []
    for line in lines:
        cells = [c.strip() for c in line.split('|')]
        if len(cells) >= 2:
            if cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
        parsed_rows.append(cells)
    if not parsed_rows:
        return None
    headers = parsed_rows[0]
    rows = []
    start_idx = 1
    if len(parsed_rows) > 1:
        is_divider = all(all(c in ('-', ':', ' ') for c in cell) for cell in parsed_rows[1] if cell)
        if is_divider:
            start_idx = 2
    for r in parsed_rows[start_idx:]:
        rows.append(r)
    return ('table', headers, rows)

def escape_xml(text):
    if not text:
        return ""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;').replace('>', '&gt;')
    return text

def append_markdown_elements_to_docx(doc, elements):
    for el in elements:
        el_type = el[0]
        if el_type == 'p':
            p = doc.add_paragraph()
            run = p.add_run(el[1])
            apply_text_formatting(run, font_name="Arial Narrow")
            set_paragraph_spacing(p)
        elif el_type == 'h':
            level = el[1]
            p = doc.add_paragraph()
            run = p.add_run(el[2])
            size = 14 if level == 1 else 12
            apply_text_formatting(run, font_name="Arial Narrow", size=size, bold=True)
            set_paragraph_spacing(p, before=12, after=6)
        elif el_type == 'li':
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(el[1])
            apply_text_formatting(run, font_name="Arial Narrow")
            set_paragraph_spacing(p, before=2, after=2)
        elif el_type == 'table':
            headers, rows = el[1], el[2]
            cols_count = max(len(headers), max((len(r) for r in rows), default=0))
            if cols_count == 0:
                continue
            table = doc.add_table(rows=1, cols=cols_count)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            for idx, h in enumerate(headers):
                if idx < len(hdr_cells):
                    hdr_cells[idx].text = h
                    for p in hdr_cells[idx].paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in p.runs:
                            apply_text_formatting(run, font_name="Arial Narrow", size=9, bold=True)
            for r in rows:
                row_cells = table.add_row().cells
                for idx, val in enumerate(r):
                    if idx < len(row_cells):
                        row_cells[idx].text = val
                        for p in row_cells[idx].paragraphs:
                            for run in p.runs:
                                apply_text_formatting(run, font_name="Arial Narrow", size=9)
            apply_three_line_table_style(table)

def append_markdown_elements_to_pdf_story(story, elements, style_normal, style_section):
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    
    style_table_header = ParagraphStyle(
        'TableHeader',
        parent=style_normal,
        fontName=style_section.fontName,
        fontSize=8,
        leading=10,
        alignment=1,
        textColor=colors.whitesmoke
    )
    style_table_cell = ParagraphStyle(
        'TableCell',
        parent=style_normal,
        fontSize=8,
        leading=10,
        alignment=0
    )
    for el in elements:
        el_type = el[0]
        if el_type == 'p':
            text_cleaned = escape_xml(el[1])
            story.append(Paragraph(text_cleaned, style_normal))
        elif el_type == 'h':
            text_cleaned = escape_xml(el[2])
            story.append(Paragraph(text_cleaned, style_section))
        elif el_type == 'li':
            text_cleaned = escape_xml(el[1])
            bullet_text = f"&bull; {text_cleaned}"
            story.append(Paragraph(bullet_text, style_normal))
        elif el_type == 'table':
            headers, rows = el[1], el[2]
            cols_count = max(len(headers), max((len(r) for r in rows), default=0))
            if cols_count == 0:
                continue
            t_data = []
            hdr_row = []
            for h in headers:
                hdr_row.append(Paragraph(escape_xml(h), style_table_header))
            while len(hdr_row) < cols_count:
                hdr_row.append(Paragraph("", style_table_header))
            t_data.append(hdr_row)
            for r in rows:
                data_row = []
                for val in r:
                    data_row.append(Paragraph(escape_xml(val), style_table_cell))
                while len(data_row) < cols_count:
                    data_row.append(Paragraph("", style_table_cell))
                t_data.append(data_row)
            available_width = 15.5 * cm
            col_width = available_width / cols_count
            t = Table(t_data, colWidths=[col_width] * cols_count)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('TOPPADDING', (0,0), (-1,-1), 4),
                ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                ('LEFTPADDING', (0,0), (-1,-1), 4),
                ('RIGHTPADDING', (0,0), (-1,-1), 4),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.5*cm))

GLOBAL_FONT_FAMILY = None
GLOBAL_LINE_SPACING = None

def get_font_name(is_article=False, specified=None):
    if GLOBAL_FONT_FAMILY is not None:
        return GLOBAL_FONT_FAMILY
    if specified is not None:
        return specified
    return "Times New Roman" if is_article else "Arial Narrow"

def get_line_spacing(is_article=False, specified=None):
    if GLOBAL_LINE_SPACING is not None:
        return GLOBAL_LINE_SPACING
    if specified is not None:
        return specified
    return 1.5 if is_article else 2.0

def get_reportlab_fonts(is_article=False):
    font_name = GLOBAL_FONT_FAMILY or ("Times New Roman" if is_article else "Helvetica")
    fn = font_name.lower()
    if "times" in fn or "roman" in fn or "serif" in fn:
        return "Times-Roman", "Times-Bold", "Times-Italic"
    elif "courier" in fn or "mono" in fn:
        return "Courier", "Courier-Bold", "Courier-Oblique"
    else:
        return "Helvetica", "Helvetica-Bold", "Helvetica-Oblique"

def set_paragraph_spacing(p, line_spacing=None, before=0, after=6):
    pf = p.paragraph_format
    pf.line_spacing = get_line_spacing(is_article=False, specified=line_spacing)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_article_paragraph_spacing(p, line_spacing=None, before=0, after=0, first_line_indent=1.27):
    pf = p.paragraph_format
    pf.line_spacing = get_line_spacing(is_article=True, specified=line_spacing)
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        pf.first_line_indent = Cm(first_line_indent)

def apply_text_formatting(run, font_name=None, size=12, bold=False, italic=False, is_article=False):
    run.font.name = get_font_name(is_article, font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def apply_section_margins(section, template_styles, default_left=3.0):
    margins = (template_styles or {}).get("margins", {})
    section.top_margin = Cm(margins.get("top", 2.5))
    section.bottom_margin = Cm(margins.get("bottom", 2.5))
    section.left_margin = Cm(margins.get("left", default_left))
    section.right_margin = Cm(margins.get("right", 2.5))

def apply_three_line_table_style(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
        
    tblBorders = OxmlElement('w:tblBorders')
    
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '8')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), 'auto')
    tblBorders.append(top)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'auto')
    tblBorders.append(bottom)
    
    for side in ['left', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'none')
        tblBorders.append(border)
        
    tblPr.append(tblBorders)
    
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        b_bottom = OxmlElement('w:bottom')
        b_bottom.set(qn('w:val'), 'single')
        b_bottom.set(qn('w:sz'), '8')
        b_bottom.set(qn('w:space'), '0')
        b_bottom.set(qn('w:color'), 'auto')
        tcBorders.append(b_bottom)
        tcPr.append(tcBorders)

# ---------------------------------------------------------------------------
# ReportLab custom canvas and flowables
# ---------------------------------------------------------------------------

class ResetPageNumberedCanvas(Flowable):
    def __init__(self):
        super(ResetPageNumberedCanvas, self).__init__()
    def draw(self):
        self.canv.first_main_page = self.canv._pageNumber

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []
        self.first_main_page = None

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        # Do not print page number on the cover (page 1)
        if self._pageNumber == 1:
            return
        self.saveState()
        self.setFont("Helvetica", 9)
        
        # Check if we have hit the main body yet
        first_main = self.first_main_page
        if first_main is not None and self._pageNumber >= first_main:
            arabic_page = self._pageNumber - first_main + 1
            page_text = f"{arabic_page}"
        else:
            page_text = to_roman(self._pageNumber)
            
        self.drawRightString(8.5 * inch - 2.5 * cm, 1.5 * cm, page_text)
        self.restoreState()

class ArticleNumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(ArticleNumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont("Times-Roman", 10)
        page_text = f"{self._pageNumber}"
        # Upper right corner page numbering
        self.drawRightString(8.5 * inch - 2.54 * cm, 11 * inch - 1.5 * cm, page_text)
        self.restoreState()

# ---------------------------------------------------------------------------
# Thesis Builders (50-page structure with Cap I, II, III, Refs, Anexos)
# ---------------------------------------------------------------------------

def build_docx(data, output_path):
    doc = Document()
    
    # ------------------ COVER PAGE (SECTION 1) ------------------
    template_styles = data.get("templateStyles", {}) or {}
    section1 = doc.sections[0]
    apply_section_margins(section1, template_styles, default_left=3.0)
    
    meta = data.get("metadata", {})
    
    # Cover layout
    structure = template_styles.get("structure", {}) or {}
    institution = (structure.get("institution") or meta.get("universidad") or "UNIVERSIDAD NACIONAL DE TRUJILLO").upper()
    faculty = (structure.get("faculty") or meta.get("facultad") or "FACULTAD DE INGENIERÍA").upper()
    school = (structure.get("school") or meta.get("escuela") or "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS").upper()
    
    school_clean = school.replace("ESCUELA PROFESIONAL DE ", "").replace("ESCUELA DE ", "").replace("CARRERA DE ", "").replace("PROGRAMA DE ", "")
    if school_clean.startswith("INGENIERÍA DE "):
        optar_titulo = "INGENIERO DE " + school_clean[14:]
    elif school_clean.startswith("INGENIERÍA "):
        optar_titulo = "INGENIERO " + school_clean[11:]
    else:
        optar_titulo = school_clean

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{institution}\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    run2 = p.add_run(f"{faculty}\n{school}\n\n")
    apply_text_formatting(run2, font_name="Arial Narrow", size=14, bold=True)
    
    logo_data = template_styles.get("logo")
    logo_inserted = False
    if logo_data and logo_data.get("base64"):
        try:
            logo_bytes = base64.b64decode(logo_data["base64"])
            ext = logo_data.get("ext", "png")
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp_img:
                tmp_img.write(logo_bytes)
                tmp_img_path = tmp_img.name
            
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(tmp_img_path, width=Inches(1.8))
            os.unlink(tmp_img_path)
            logo_inserted = True
        except Exception as img_err:
            print(f"Error inserting logo in docx: {img_err}", file=sys.stderr)
            
    p_spacing = doc.add_paragraph()
    p_spacing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing_text = "\n\n" if logo_inserted else "\n\n\n\n"
    run_space = p_spacing.add_run(spacing_text)
    apply_text_formatting(run_space, font_name="Arial Narrow", size=12)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"\"{meta.get('titulo_proyecto', '').upper()}\"\n\n")
    apply_text_formatting(run_title, font_name="Arial Narrow", size=14, bold=True)
    
    p_spacing2 = doc.add_paragraph()
    p_spacing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing_text2 = "\n" if logo_inserted else "\n\n\n"
    run_space2 = p_spacing2.add_run(spacing_text2)
    apply_text_formatting(run_space2, font_name="Arial Narrow", size=12)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"PROYECTO DE TESIS PARA OPTAR EL TÍTULO PROFESIONAL DE\n{optar_titulo}\n\n\n")
    apply_text_formatting(run_sub, font_name="Arial Narrow", size=12, bold=True)
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_author.add_run(f"Autor: {meta.get('nombre_autor', '')}\nAsesor: Dr. {meta.get('nombre_asesor', '')}\n\n")
    apply_text_formatting(run_auth, font_name="Arial Narrow", size=12)
    
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run(f"Línea de Investigación: {meta.get('linea_investigacion', '')}\n\n\n\n")
    apply_text_formatting(run_line, font_name="Arial Narrow", size=11, italic=True)
    
    p_city = doc.add_paragraph()
    p_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_city = p_city.add_run(f"{meta.get('ciudad', '')} - Perú\n{meta.get('anio', '')}")
    apply_text_formatting(run_city, font_name="Arial Narrow", size=12, bold=True)
    
    # ------------------ PRELIMINARES (SECTION 2) ------------------
    section2 = doc.add_section()
    apply_section_margins(section2, template_styles, default_left=3.0)
    section2.footer.is_linked_to_previous = False
    configure_page_numbering(section2, fmt="romanLower", start=2)
    p_footer2 = section2.footer.paragraphs[0]
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number_to_footer(p_footer2)

    # Jurado Dictaminador
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nJURADO DICTAMINADOR\n\n\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    
    jurados = [
        ("Presidente", "Dr. Roberto Carlos Medina"),
        ("Secretario", "Dr. Julio César Alvarez"),
        ("Vocal (Asesor)", f"Dr. {meta.get('nombre_asesor', '')}")
    ]
    for role, name in jurados:
        p_j = doc.add_paragraph()
        p_j.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_j.paragraph_format.space_before = Pt(30)
        run_line = p_j.add_run("_______________________________\n")
        apply_text_formatting(run_line, font_name="Arial Narrow", size=12)
        run_name = p_j.add_run(f"{name}\n")
        apply_text_formatting(run_name, font_name="Arial Narrow", size=12, bold=True)
        run_role = p_j.add_run(role)
        apply_text_formatting(run_role, font_name="Arial Narrow", size=11, italic=True)
        
    doc.add_page_break()

    # Dedicatoria & Agradecimientos
    prelims = data.get("preliminares", {})
    if prelims.get("dedicatoria"):
        p = doc.add_paragraph()
        run = p.add_run("DEDICATORIA\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_ded = doc.add_paragraph()
        p_ded.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ded.paragraph_format.left_indent = Cm(5.0)
        run_ded = p_ded.add_run(prelims.get("dedicatoria"))
        apply_text_formatting(run_ded, font_name="Arial Narrow", italic=True)
        set_paragraph_spacing(p_ded)
        doc.add_page_break()
        
    if prelims.get("agradecimientos"):
        p = doc.add_paragraph()
        run = p.add_run("AGRADECIMIENTOS\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_agr = doc.add_paragraph()
        run_agr = p_agr.add_run(prelims.get("agradecimientos"))
        apply_text_formatting(run_agr, font_name="Arial Narrow")
        set_paragraph_spacing(p_agr)
        doc.add_page_break()

    # Índice General
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE GENERAL\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    
    for item in prelims.get("indice_general", []):
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.tab_stops.add_tab_stop(Cm(14.0), alignment=2)
        run_item = p_toc.add_run(f"{item}")
        apply_text_formatting(run_item, font_name="Arial Narrow")
        set_paragraph_spacing(p_toc, line_spacing=1.2, before=0, after=2)
        
    doc.add_page_break()

    # Índice de tablas
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE DE TABLAS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    for idx, item in enumerate(prelims.get("indice_tablas", [])):
        p_tbl = doc.add_paragraph()
        run_tbl = p_tbl.add_run(item)
        apply_text_formatting(run_tbl, font_name="Arial Narrow")
        set_paragraph_spacing(p_tbl, line_spacing=1.2, before=0, after=2)
    doc.add_page_break()

    # Índice de figuras
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE DE FIGURAS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    for idx, item in enumerate(prelims.get("indice_figuras", [])):
        p_fig = doc.add_paragraph()
        run_fig = p_fig.add_run(item)
        apply_text_formatting(run_fig, font_name="Arial Narrow")
        set_paragraph_spacing(p_fig, line_spacing=1.2, before=0, after=2)
    doc.add_page_break()

    # Presentación
    if prelims.get("presentacion"):
        p = doc.add_paragraph()
        run = p.add_run("PRESENTACIÓN\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_pres = doc.add_paragraph()
        run_pres = p_pres.add_run(prelims.get("presentacion"))
        apply_text_formatting(run_pres, font_name="Arial Narrow")
        set_paragraph_spacing(p_pres)
        doc.add_page_break()
        
    # Resumen
    if prelims.get("resumen"):
        p = doc.add_paragraph()
        run = p.add_run("RESUMEN\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_res = doc.add_paragraph()
        run_res = p_res.add_run(prelims.get("resumen"))
        apply_text_formatting(run_res, font_name="Arial Narrow")
        set_paragraph_spacing(p_res)
        
        if prelims.get("palabras_clave"):
            p_kw = doc.add_paragraph()
            run_lbl = p_kw.add_run("Palabras clave: ")
            apply_text_formatting(run_lbl, font_name="Arial Narrow", bold=True)
            run_val = p_kw.add_run(prelims.get("palabras_clave"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_kw, before=12)
            
        doc.add_page_break()
        
    # Abstract
    if prelims.get("abstract"):
        p = doc.add_paragraph()
        run = p.add_run("ABSTRACT\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_abs = doc.add_paragraph()
        run_abs = p_abs.add_run(prelims.get("abstract"))
        apply_text_formatting(run_abs, font_name="Arial Narrow")
        set_paragraph_spacing(p_abs)
        
        if prelims.get("keywords"):
            p_kw = doc.add_paragraph()
            run_lbl = p_kw.add_run("Keywords: ")
            apply_text_formatting(run_lbl, font_name="Arial Narrow", bold=True)
            run_val = p_kw.add_run(prelims.get("keywords"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_kw, before=12)
            
    # ------------------ MAIN BODY (SECTION 3) ------------------
    section3 = doc.add_section()
    apply_section_margins(section3, template_styles, default_left=3.0)
    section3.footer.is_linked_to_previous = False
    configure_page_numbering(section3, fmt="decimal", start=1)
    p_footer3 = section3.footer.paragraphs[0]
    p_footer3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number_to_footer(p_footer3)

    # CAPÍTULO I: INTRODUCCIÓN
    c1 = data.get("capitulo1", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO I: INTRODUCCIÓN\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    sections_c1 = c1.get("secciones", [])
    if sections_c1 and isinstance(sections_c1, list):
        for sec in sections_c1:
            title = sec.get("titulo", "")
            text = sec.get("contenido", "")
            if title and text:
                p_title = doc.add_paragraph()
                run_t = p_title.add_run(title)
                apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
                set_paragraph_spacing(p_title, before=12, after=6)
                
                p_text = doc.add_paragraph()
                run_txt = p_text.add_run(get_clean_text(text))
                apply_text_formatting(run_txt, font_name="Arial Narrow")
                set_paragraph_spacing(p_text)
    else:
        sections_c1_default = [
            ("1.1 Realidad problemática", c1.get("realidad_problematica")),
            ("1.2 Antecedentes de la investigación", c1.get("antecedentes")),
            ("1.3 Marco teórico", c1.get("marco_teorico")),
            ("1.4 Metodologías alternativas", c1.get("metodologias_alternativas")),
            ("1.5 Justificación de la investigación", c1.get("justificacion")),
            ("1.6 Formulación del problema", c1.get("formulacion_problema")),
            ("1.7 Hipótesis", c1.get("hipotesis")),
            ("1.8 Objetivos", c1.get("objetivos")),
            ("1.9 Limitaciones del estudio", c1.get("limitaciones"))
        ]
        for title, text in sections_c1_default:
            if text:
                p_title = doc.add_paragraph()
                run_t = p_title.add_run(title)
                apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
                set_paragraph_spacing(p_title, before=12, after=6)
                
                p_text = doc.add_paragraph()
                run_txt = p_text.add_run(get_clean_text(text))
                apply_text_formatting(run_txt, font_name="Arial Narrow")
                set_paragraph_spacing(p_text)
                
    doc.add_page_break()

    # CAPÍTULO II: MÉTODO
    c2 = data.get("capitulo2", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO II: MÉTODO\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    sections_c2 = c2.get("secciones", [])
    if sections_c2 and isinstance(sections_c2, list):
        variables_rendered = False
        for sec in sections_c2:
            title = sec.get("titulo", "")
            text = sec.get("contenido", "")
            if title and text:
                p_title = doc.add_paragraph()
                run_t = p_title.add_run(title)
                apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
                set_paragraph_spacing(p_title, before=12, after=6)
                
                p_text = doc.add_paragraph()
                run_txt = p_text.add_run(get_clean_text(text))
                apply_text_formatting(run_txt, font_name="Arial Narrow")
                set_paragraph_spacing(p_text)
                
                # Check if this section relates to variables & operacionalizacion
                title_lower = title.lower()
                if ("variable" in title_lower or "operacionaliz" in title_lower) and not variables_rendered:
                    vars_data = c2.get("variables", {})
                    op_table_data = vars_data.get("operacionalizacion_tabla", []) if vars_data else []
                    if op_table_data:
                        p_lbl = doc.add_paragraph()
                        run_lbl = p_lbl.add_run("Tabla 1. Matriz de Operacionalización de Variables")
                        apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                        
                        table = doc.add_table(rows=1, cols=6)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        hdr_cells = table.rows[0].cells
                        hdr_titles = ['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']
                        for idx, col_title in enumerate(hdr_titles):
                            hdr_cells[idx].text = col_title
                            apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        for r in op_table_data:
                            row_cells = table.add_row().cells
                            row_cells[0].text = r.get("variable", "")
                            row_cells[1].text = r.get("definicion_conceptual", "")
                            row_cells[2].text = r.get("definicion_operacional", "")
                            row_cells[3].text = r.get("dimensiones", "")
                            row_cells[4].text = r.get("indicadores", "")
                            row_cells[5].text = r.get("escala_medicion", "")
                            for cell in row_cells:
                                apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                        variables_rendered = True

        # Append variables table if it wasn't rendered
        if not variables_rendered:
            vars_data = c2.get("variables", {})
            op_table_data = vars_data.get("operacionalizacion_tabla", []) if vars_data else []
            if op_table_data:
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("Variables y Operacionalización")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                
                p_lbl = doc.add_paragraph()
                run_lbl = p_lbl.add_run("Tabla 1. Matriz de Operacionalización de Variables")
                apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                
                table = doc.add_table(rows=1, cols=6)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                hdr_titles = ['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']
                for idx, col_title in enumerate(hdr_titles):
                    hdr_cells[idx].text = col_title
                    apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                for r in op_table_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r.get("variable", "")
                    row_cells[1].text = r.get("definicion_conceptual", "")
                    row_cells[2].text = r.get("definicion_operacional", "")
                    row_cells[3].text = r.get("dimensiones", "")
                    row_cells[4].text = r.get("indicadores", "")
                    row_cells[5].text = r.get("escala_medicion", "")
                    for cell in row_cells:
                        apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
    else:
        # Fallback default sequence
        sections_c2 = [
            ("2.1 Tipo de investigación", c2.get("tipo_investigacion")),
            ("2.2 Nivel de investigación", c2.get("nivel_investigacion")),
            ("2.3 Diseño de investigación", c2.get("diseno_investigacion"))
        ]
        for title, text in sections_c2:
            if text:
                p_title = doc.add_paragraph()
                run_t = p_title.add_run(title)
                apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
                set_paragraph_spacing(p_title, before=12, after=6)
                p_text = doc.add_paragraph()
                run_txt = p_text.add_run(text)
                apply_text_formatting(run_txt, font_name="Arial Narrow")
                set_paragraph_spacing(p_text)

        # Población y muestra
        p_pm = c2.get("poblacion_muestra", {})
        if p_pm:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("2.4 Población, muestra y muestreo")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            for label, key in [
                ("2.4.1 Población", "poblacion"),
                ("2.4.2 Muestra", "muestra"),
                ("2.4.3 Muestreo", "muestreo")
            ]:
                val = p_pm.get(key)
                if val:
                    p_sub = doc.add_paragraph()
                    run_sub = p_sub.add_run(label)
                    apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                    set_paragraph_spacing(p_sub, before=8, after=4)
                    
                    p_val = doc.add_paragraph()
                    run_val = p_val.add_run(val)
                    apply_text_formatting(run_val, font_name="Arial Narrow")
                    set_paragraph_spacing(p_val)

        # Variables y matriz
        vars_data = c2.get("variables", {})
        if vars_data:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("2.5 Variables")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            if vars_data.get("tipo"):
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("2.5.1 Tipo de Variables")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                p_val = doc.add_paragraph()
                run_val = p_val.add_run(vars_data.get("tipo"))
                apply_text_formatting(run_val, font_name="Arial Narrow")
                set_paragraph_spacing(p_val)
                
            op_table_data = vars_data.get("operacionalizacion_tabla", [])
            if op_table_data:
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("2.5.2 Operacionalización de Variables")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                
                p_lbl = doc.add_paragraph()
                run_lbl = p_lbl.add_run("Tabla 1. Matriz de Operacionalización de Variables")
                apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                
                table = doc.add_table(rows=1, cols=6)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                hdr_titles = ['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']
                for idx, col_title in enumerate(hdr_titles):
                    hdr_cells[idx].text = col_title
                    apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                for r in op_table_data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r.get("variable", "")
                    row_cells[1].text = r.get("definicion_conceptual", "")
                    row_cells[2].text = r.get("definicion_operacional", "")
                    row_cells[3].text = r.get("dimensiones", "")
                    row_cells[4].text = r.get("indicadores", "")
                    row_cells[5].text = r.get("escala_medicion", "")
                    for cell in row_cells:
                        apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)

        # Técnicas e instrumentos
        ti_data = c2.get("tecnicas_instrumentos", {})
        if ti_data:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("2.6 Técnicas e instrumentos de recolección de datos")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            if ti_data.get("descripcion"):
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("2.6.1 Descripción de técnicas e instrumentos")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                p_val = doc.add_paragraph()
                run_val = p_val.add_run(ti_data.get("descripcion"))
                apply_text_formatting(run_val, font_name="Arial Narrow")
                set_paragraph_spacing(p_val)
                
            if ti_data.get("validacion_confiabilidad"):
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("2.6.2 Validación y confiabilidad")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                p_val = doc.add_paragraph()
                run_val = p_val.add_run(ti_data.get("validacion_confiabilidad"))
                apply_text_formatting(run_val, font_name="Arial Narrow")
                set_paragraph_spacing(p_val)

        sections_c2_end = [
            ("2.7 Método de análisis de datos", c2.get("metodo_analisis")),
            ("2.8 Procedimiento", c2.get("procedimiento")),
            ("2.9 Consideraciones éticas", c2.get("consideraciones_eticas"))
        ]
        for title, text in sections_c2_end:
            if text:
                p_title = doc.add_paragraph()
                run_t = p_title.add_run(title)
                apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
                set_paragraph_spacing(p_title, before=12, after=6)
                p_text = doc.add_paragraph()
                run_txt = p_text.add_run(text)
                apply_text_formatting(run_txt, font_name="Arial Narrow")
                set_paragraph_spacing(p_text)

    doc.add_page_break()

    # CAPÍTULO III: ASPECTOS ADMINISTRATIVOS
    c3 = data.get("capitulo3", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO III: ASPECTOS ADMINISTRATIVOS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    sections_c3 = c3.get("secciones", [])
    if sections_c3 and isinstance(sections_c3, list):
        budget_rendered = False
        cronograma_rendered = False
        for sec in sections_c3:
            title = sec.get("titulo", "")
            text = sec.get("contenido", "")
            if title and text:
                p_title = doc.add_paragraph()
                run_t = p_title.add_run(title)
                apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
                set_paragraph_spacing(p_title, before=12, after=6)
                
                p_text = doc.add_paragraph()
                run_txt = p_text.add_run(get_clean_text(text))
                apply_text_formatting(run_txt, font_name="Arial Narrow")
                set_paragraph_spacing(p_text)
                
                title_lower = title.lower()
                
                # Check for budget table
                if ("presupuesto" in title_lower or "costo" in title_lower or "financiamiento" in title_lower) and not budget_rendered:
                    pres_table = c3.get("presupuesto_tabla", [])
                    if pres_table:
                        p_tbl_lbl = doc.add_paragraph()
                        run_lbl = p_tbl_lbl.add_run("Tabla 2. Presupuesto detallado del proyecto de tesis")
                        apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                        set_paragraph_spacing(p_tbl_lbl, before=6, after=6)
                        
                        table = doc.add_table(rows=1, cols=6)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        hdr_cells = table.rows[0].cells
                        hdr_titles = ['Categoría', 'Recurso', 'Unidad', 'Costo Unitario', 'Cantidad', 'Total']
                        for idx, col_title in enumerate(hdr_titles):
                            hdr_cells[idx].text = col_title
                            apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        total_presupuesto = 0
                        for r in pres_table:
                            row_cells = table.add_row().cells
                            row_cells[0].text = r.get("categoria", "")
                            row_cells[1].text = r.get("recurso", "")
                            row_cells[2].text = r.get("unidad", "")
                            row_cells[3].text = f"S/. {r.get('costo_unitario', 0)}"
                            row_cells[4].text = str(r.get('cantidad', 0))
                            
                            costo_total = r.get('costo_total', 0)
                            if costo_total == 0:
                                costo_total = r.get('costo_unitario', 0) * r.get('cantidad', 0)
                            total_presupuesto += costo_total
                            row_cells[5].text = f"S/. {costo_total}"
                            
                            for cell in row_cells:
                                apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                                
                        row_cells = table.add_row().cells
                        row_cells[0].text = "TOTAL PRESUPUESTO"
                        row_cells[5].text = f"S/. {total_presupuesto}"
                        apply_text_formatting(row_cells[0].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                        apply_text_formatting(row_cells[5].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                        budget_rendered = True
                
                # Check for cronograma table
                if ("cronograma" in title_lower or "actividad" in title_lower or "gantt" in title_lower or "ejecución" in title_lower or "ejecucion" in title_lower) and not cronograma_rendered:
                    crono_data = c3.get("cronograma", {})
                    crono_table = crono_data.get("cronograma_tabla", []) if crono_data else []
                    if crono_table:
                        p_tbl_lbl = doc.add_paragraph()
                        run_lbl = p_tbl_lbl.add_run("Tabla 3. Cronograma Gantt de actividades")
                        apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                        set_paragraph_spacing(p_tbl_lbl, before=6, after=6)
                        
                        table = doc.add_table(rows=1, cols=7)
                        table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        hdr_cells = table.rows[0].cells
                        hdr_titles = ['Actividad', 'Mes 1', 'Mes 2', 'Mes 3', 'Mes 4', 'Mes 5', 'Mes 6']
                        for idx, col_title in enumerate(hdr_titles):
                            hdr_cells[idx].text = col_title
                            apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                            hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        for r in crono_table:
                            row_cells = table.add_row().cells
                            row_cells[0].text = r.get("actividad", "")
                            row_cells[1].text = r.get("mes_1", "")
                            row_cells[2].text = r.get("mes_2", "")
                            row_cells[3].text = r.get("mes_3", "")
                            row_cells[4].text = r.get("mes_4", "")
                            row_cells[5].text = r.get("mes_5", "")
                            row_cells[6].text = r.get("mes_6", "")
                            for idx, cell in enumerate(row_cells):
                                apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                                if idx > 0:
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cronograma_rendered = True

        # Append tables if not rendered
        if not budget_rendered:
            pres_table = c3.get("presupuesto_tabla", [])
            if pres_table:
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("Presupuesto")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                
                p_tbl_lbl = doc.add_paragraph()
                run_lbl = p_tbl_lbl.add_run("Tabla 2. Presupuesto detallado del proyecto de tesis")
                apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                
                table = doc.add_table(rows=1, cols=6)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                hdr_titles = ['Categoría', 'Recurso', 'Unidad', 'Costo Unitario', 'Cantidad', 'Total']
                for idx, col_title in enumerate(hdr_titles):
                    hdr_cells[idx].text = col_title
                    apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                total_presupuesto = 0
                for r in pres_table:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r.get("categoria", "")
                    row_cells[1].text = r.get("recurso", "")
                    row_cells[2].text = r.get("unidad", "")
                    row_cells[3].text = f"S/. {r.get('costo_unitario', 0)}"
                    row_cells[4].text = str(r.get('cantidad', 0))
                    
                    costo_total = r.get('costo_total', 0)
                    if costo_total == 0:
                        costo_total = r.get('costo_unitario', 0) * r.get('cantidad', 0)
                    total_presupuesto += costo_total
                    row_cells[5].text = f"S/. {costo_total}"
                    
                    for cell in row_cells:
                        apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                        
                row_cells = table.add_row().cells
                row_cells[0].text = "TOTAL PRESUPUESTO"
                row_cells[5].text = f"S/. {total_presupuesto}"
                apply_text_formatting(row_cells[0].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                apply_text_formatting(row_cells[5].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)

        if not cronograma_rendered:
            crono_data = c3.get("cronograma", {})
            crono_table = crono_data.get("cronograma_tabla", []) if crono_data else []
            if crono_table:
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run("Cronograma de Actividades")
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                
                p_tbl_lbl = doc.add_paragraph()
                run_lbl = p_tbl_lbl.add_run("Tabla 3. Cronograma Gantt de actividades")
                apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                
                table = doc.add_table(rows=1, cols=7)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                hdr_titles = ['Actividad', 'Mes 1', 'Mes 2', 'Mes 3', 'Mes 4', 'Mes 5', 'Mes 6']
                for idx, col_title in enumerate(hdr_titles):
                    hdr_cells[idx].text = col_title
                    apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                for r in crono_table:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r.get("actividad", "")
                    row_cells[1].text = r.get("mes_1", "")
                    row_cells[2].text = r.get("mes_2", "")
                    row_cells[3].text = r.get("mes_3", "")
                    row_cells[4].text = r.get("mes_4", "")
                    row_cells[5].text = r.get("mes_5", "")
                    row_cells[6].text = r.get("mes_6", "")
                    for idx, cell in enumerate(row_cells):
                        apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                        if idx > 0:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Fallback default sequence
        rec_data = c3.get("recursos", {})
        if rec_data:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("3.1 Recursos")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            for label, key in [
                ("3.1.1 Personal", "personal"),
                ("3.1.2 Bienes", "bienes"),
                ("3.1.3 Viajes", "viajes"),
                ("3.1.4 Servicios", "servicios"),
                ("3.1.5 Tecnológicos", "tecnologicos")
            ]:
                val = rec_data.get(key)
                if val:
                    p_sub = doc.add_paragraph()
                    run_st = p_sub.add_run(label)
                    apply_text_formatting(run_st, font_name="Arial Narrow", size=11, bold=True)
                    set_paragraph_spacing(p_sub, before=8, after=4)
                    p_val = doc.add_paragraph()
                    run_val = p_val.add_run(val)
                    apply_text_formatting(run_val, font_name="Arial Narrow")
                    set_paragraph_spacing(p_val)
                    
        pres_table = c3.get("presupuesto_tabla", [])
        if pres_table:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("3.2 Presupuesto")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            p_tbl_lbl = doc.add_paragraph()
            run_lbl = p_tbl_lbl.add_run("Tabla 2. Presupuesto detallado del proyecto de tesis")
            apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_tbl_lbl, before=6, after=6)
            
            table = doc.add_table(rows=1, cols=6)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_titles = ['Categoría', 'Recurso', 'Unidad', 'Costo Unitario', 'Cantidad', 'Total']
            for idx, col_title in enumerate(hdr_titles):
                hdr_cells[idx].text = col_title
                apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            total_presupuesto = 0
            for r in pres_table:
                row_cells = table.add_row().cells
                row_cells[0].text = r.get("categoria", "")
                row_cells[1].text = r.get("recurso", "")
                row_cells[2].text = r.get("unidad", "")
                row_cells[3].text = f"S/. {r.get('costo_unitario', 0)}"
                row_cells[4].text = str(r.get('cantidad', 0))
                
                costo_total = r.get('costo_total', 0)
                if costo_total == 0:
                    costo_total = r.get('costo_unitario', 0) * r.get('cantidad', 0)
                total_presupuesto += costo_total
                row_cells[5].text = f"S/. {costo_total}"
                
                for cell in row_cells:
                    apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                    
            # Total row
            row_cells = table.add_row().cells
            row_cells[0].text = "TOTAL PRESUPUESTO"
            row_cells[5].text = f"S/. {total_presupuesto}"
            apply_text_formatting(row_cells[0].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
            apply_text_formatting(row_cells[5].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)

        fin_text = c3.get("financiamiento")
        if fin_text:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("3.3 Financiamiento")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            p_text = doc.add_paragraph()
            run_txt = p_text.add_run(fin_text)
            apply_text_formatting(run_txt, font_name="Arial Narrow")
            set_paragraph_spacing(p_text)
            
        crono_data = c3.get("cronograma", {})
        if crono_data:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run("3.4 Cronograma de ejecución")
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            if crono_data.get("periodo"):
                p_sub = doc.add_paragraph()
                run_st = p_sub.add_run(f"3.4.1 Período: {crono_data.get('periodo')}")
                apply_text_formatting(run_st, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                
            p_sub2 = doc.add_paragraph()
            run_st2 = p_sub2.add_run("3.4.2 Cronograma de Actividades")
            apply_text_formatting(run_st2, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_sub2, before=8, after=4)
            
            crono_table = crono_data.get("cronograma_tabla", [])
            if crono_table:
                p_tbl_lbl = doc.add_paragraph()
                run_lbl = p_tbl_lbl.add_run("Tabla 3. Cronograma Gantt de actividades")
                apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_tbl_lbl, before=6, after=6)
                
                table = doc.add_table(rows=1, cols=7)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                hdr_cells = table.rows[0].cells
                hdr_titles = ['Actividad', 'Mes 1', 'Mes 2', 'Mes 3', 'Mes 4', 'Mes 5', 'Mes 6']
                for idx, col_title in enumerate(hdr_titles):
                    hdr_cells[idx].text = col_title
                    apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                    hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                for r in crono_table:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r.get("actividad", "")
                    row_cells[1].text = r.get("mes_1", "")
                    row_cells[2].text = r.get("mes_2", "")
                    row_cells[3].text = r.get("mes_3", "")
                    row_cells[4].text = r.get("mes_4", "")
                    row_cells[5].text = r.get("mes_5", "")
                    row_cells[6].text = r.get("mes_6", "")
                    for idx, cell in enumerate(row_cells):
                        apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
                        if idx > 0:
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # REFERENCIAS BIBLIOGRÁFICAS
    p = doc.add_paragraph()
    run = p.add_run("REFERENCIAS BIBLIOGRÁFICAS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    refs = data.get("referencias", [])
    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Cm(1.27)
        p_ref.paragraph_format.first_line_indent = Cm(-1.27)
        run_ref = p_ref.add_run(ref)
        apply_text_formatting(run_ref, font_name="Arial Narrow", size=11)
        set_paragraph_spacing(p_ref, line_spacing=1.5, before=0, after=6)
        
    doc.add_page_break()

    # ANEXOS OBLIGATORIOS
    p = doc.add_paragraph()
    run = p.add_run("ANEXOS OBLIGATORIOS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    anexos = data.get("anexos", {})
    sections_an = [
        ("Anexo 1. Matriz de operacionalización de variables", anexos.get("anexo_1")),
        ("Anexo 2. Matriz de consistencia", anexos.get("anexo_2")),
        ("Anexo 3. Diagrama de Ishikawa", anexos.get("anexo_3")),
        ("Anexo 4. Árbol de problemas", anexos.get("anexo_4")),
        ("Anexo 5. Árbol de objetivos", anexos.get("anexo_5")),
        ("Anexo 6. Instrumentos de recolección de datos", anexos.get("anexo_6")),
        ("Anexo 7. Constancia de aplicación de instrumentos", anexos.get("anexo_7")),
        ("Anexo 8. Declaración de originalidad y conformidad", anexos.get("anexo_8"))
    ]
    for title, text in sections_an:
        if text:
            p_t = doc.add_paragraph()
            run_t = p_t.add_run(title)
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_t, before=12, after=6)
            
            elements = parse_markdown_to_elements(text)
            append_markdown_elements_to_docx(doc, elements)

    doc.save(output_path)


def build_pdf(data, output_path):
    template_styles = data.get("templateStyles", {}) or {}
    temp_files_to_clean = []
    margins = template_styles.get("margins", {})
    
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=margins.get("left", 3.0)*cm,
        rightMargin=margins.get("right", 2.5)*cm,
        topMargin=margins.get("top", 2.5)*cm,
        bottomMargin=margins.get("bottom", 2.5)*cm
    )
    
    styles = getSampleStyleSheet()
    
    font_reg, font_bold, font_italic = get_reportlab_fonts(is_article=False)
    line_sp = get_line_spacing(is_article=False)
    
    style_normal = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontName=font_reg,
        fontSize=11,
        leading=int(11 * line_sp * 1.2),
        alignment=4,   # Justified
        spaceAfter=6
    )
    
    style_title = ParagraphStyle(
        'CustomTitle',
        parent=styles['Normal'],
        fontName=font_bold,
        fontSize=14,
        leading=20,
        alignment=1,
        spaceAfter=15
    )
    
    style_header = ParagraphStyle(
        'CustomHeader',
        parent=styles['Normal'],
        fontName=font_bold,
        fontSize=15,
        leading=22,
        alignment=1,
        spaceAfter=20
    )
    
    style_section = ParagraphStyle(
        'CustomSection',
        parent=styles['Normal'],
        fontName=font_bold,
        fontSize=12,
        leading=18,
        alignment=0,
        spaceBefore=12,
        spaceAfter=6
    )
    
    style_ref = ParagraphStyle(
        'CustomRef',
        parent=styles['Normal'],
        fontName=font_reg,
        fontSize=10,
        leading=15,
        alignment=4,
        leftIndent=1.27*cm,
        firstLineIndent=-1.27*cm,
        spaceAfter=6
    )

    story = []
    meta = data.get("metadata", {})
    
    # Cover Page
    structure = template_styles.get("structure", {}) or {}
    institution = (structure.get("institution") or meta.get("universidad") or "UNIVERSIDAD NACIONAL DE TRUJILLO").upper()
    faculty = (structure.get("faculty") or meta.get("facultad") or "FACULTAD DE INGENIERÍA").upper()
    school = (structure.get("school") or meta.get("escuela") or "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS").upper()
    
    school_clean = school.replace("ESCUELA PROFESIONAL DE ", "").replace("ESCUELA DE ", "").replace("CARRERA DE ", "").replace("PROGRAMA DE ", "")
    if school_clean.startswith("INGENIERÍA DE "):
        optar_titulo = "INGENIERO DE " + school_clean[14:]
    elif school_clean.startswith("INGENIERÍA "):
        optar_titulo = "INGENIERO " + school_clean[11:]
    else:
        optar_titulo = school_clean

    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(institution, style_title))
    story.append(Paragraph(f"{faculty}<br/>{school}", style_title))
    story.append(Spacer(1, 0.5*cm))
    
    logo_data = template_styles.get("logo")
    logo_inserted = False
    if logo_data and logo_data.get("base64"):
        try:
            logo_bytes = base64.b64decode(logo_data["base64"])
            ext = logo_data.get("ext", "png")
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp_img:
                tmp_img.write(logo_bytes)
                tmp_img_path = tmp_img.name
            
            temp_files_to_clean.append(tmp_img_path)
            from reportlab.platypus import Image
            logo_w = 1.8 * inch
            story.append(Image(tmp_img_path, width=logo_w, height=logo_w))
            logo_inserted = True
        except Exception as img_err:
            print(f"Error inserting logo in PDF: {img_err}", file=sys.stderr)
            
    if logo_inserted:
        story.append(Spacer(1, 0.8*cm))
    else:
        story.append(Spacer(1, 1.5*cm))
        
    story.append(Paragraph(f"\"{meta.get('titulo_proyecto', '').upper()}\"", style_header))
    
    if logo_inserted:
        story.append(Spacer(1, 0.8*cm))
    else:
        story.append(Spacer(1, 1.5*cm))
        
    story.append(Paragraph(f"PROYECTO DE TESIS PARA OPTAR EL TÍTULO PROFESIONAL DE<br/>{optar_titulo}", style_title))
    
    if logo_inserted:
        story.append(Spacer(1, 0.8*cm))
    else:
        story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph(f"Autor: {meta.get('nombre_autor', '')}<br/>Asesor: Dr. {meta.get('nombre_asesor', '')}", style_normal))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(f"Línea de Investigación: {meta.get('linea_investigacion', '')}", style_normal))
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(f"{meta.get('ciudad', '')} - Perú<br/>{meta.get('anio', '')}", style_title))
    story.append(PageBreak())

    # Jurado
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("JURADO DICTAMINADOR", style_title))
    story.append(Spacer(1, 1.5*cm))
    
    jurados = [
        ("Presidente", "Dr. Roberto Carlos Medina"),
        ("Secretario", "Dr. Julio César Alvarez"),
        ("Vocal (Asesor)", f"Dr. {meta.get('nombre_asesor', '')}")
    ]
    for role, name in jurados:
        story.append(Spacer(1, 1.5*cm))
        story.append(Paragraph("_______________________________", style_title))
        story.append(Paragraph(f"<b>{name}</b><br/><i>{role}</i>", style_title))
    story.append(PageBreak())

    # Dedicatoria & Agradecimientos
    prelims = data.get("preliminares", {})
    if prelims.get("dedicatoria"):
        story.append(Paragraph("DEDICATORIA", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(f"<i>{prelims.get('dedicatoria')}</i>", style_normal))
        story.append(PageBreak())
        
    if prelims.get("agradecimientos"):
        story.append(Paragraph("AGRADECIMIENTOS", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("agradecimientos"), style_normal))
        story.append(PageBreak())

    # TOC
    story.append(Paragraph("ÍNDICE GENERAL", style_title))
    story.append(Spacer(1, 0.5*cm))
    for item in prelims.get("indice_general", []):
        story.append(Paragraph(f"{item}", style_normal))
    story.append(PageBreak())

    # Índice de tablas
    story.append(Paragraph("ÍNDICE DE TABLAS", style_title))
    story.append(Spacer(1, 0.5*cm))
    for item in prelims.get("indice_tablas", []):
        story.append(Paragraph(item, style_normal))
    story.append(PageBreak())

    # Índice de figuras
    story.append(Paragraph("ÍNDICE DE FIGURAS", style_title))
    story.append(Spacer(1, 0.5*cm))
    for item in prelims.get("indice_figuras", []):
        story.append(Paragraph(item, style_normal))
    story.append(PageBreak())

    # Presentación
    if prelims.get("presentacion"):
        story.append(Paragraph("PRESENTACIÓN", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("presentacion"), style_normal))
        story.append(PageBreak())
        
    # Resumen y Abstract
    if prelims.get("resumen"):
        story.append(Paragraph("RESUMEN", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("resumen"), style_normal))
        if prelims.get("palabras_clave"):
            story.append(Paragraph(f"<b>Palabras clave:</b> {prelims.get('palabras_clave')}", style_normal))
        story.append(PageBreak())
        
    if prelims.get("abstract"):
        story.append(Paragraph("ABSTRACT", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("abstract"), style_normal))
        if prelims.get("keywords"):
            story.append(Paragraph(f"<b>Keywords:</b> {prelims.get('keywords')}", style_normal))
        story.append(PageBreak())

    # --- Start of Main Body (Roman to Arabic transition) ---
    story.append(ResetPageNumberedCanvas())

    # Capítulo I
    c1 = data.get("capitulo1", {})
    story.append(Paragraph("CAPÍTULO I: INTRODUCCIÓN", style_header))
    story.append(Spacer(1, 1*cm))
    
    sections_c1 = c1.get("secciones", [])
    if sections_c1 and isinstance(sections_c1, list):
        for sec in sections_c1:
            title = sec.get("titulo", "")
            text = sec.get("contenido", "")
            if title and text:
                story.append(Paragraph(title, style_section))
                story.append(Paragraph(get_clean_html_text(text), style_normal))
    else:
        sections = [
            ("1.1 Realidad problemática", c1.get("realidad_problematica")),
            ("1.2 Antecedentes de la investigación", c1.get("antecedentes")),
            ("1.3 Marco teórico", c1.get("marco_teorico")),
            ("1.4 Metodologías alternativas", c1.get("metodologias_alternativas")),
            ("1.5 Justificación de la investigación", c1.get("justificacion")),
            ("1.6 Formulación del problema", c1.get("formulacion_problema")),
            ("1.7 Hipótesis", c1.get("hipotesis")),
            ("1.8 Objetivos", c1.get("objetivos")),
            ("1.9 Limitaciones del estudio", c1.get("limitaciones"))
        ]
        for title, text in sections:
            if text:
                story.append(Paragraph(title, style_section))
                story.append(Paragraph(get_clean_html_text(text), style_normal))
    story.append(PageBreak())

    # Capítulo II
    c2 = data.get("capitulo2", {})
    story.append(Paragraph("CAPÍTULO II: MÉTODO", style_header))
    story.append(Spacer(1, 1*cm))
    
    sections_c2 = c2.get("secciones", [])
    if sections_c2 and isinstance(sections_c2, list):
        variables_rendered = False
        for sec in sections_c2:
            title = sec.get("titulo", "")
            text = sec.get("contenido", "")
            if title and text:
                story.append(Paragraph(title, style_section))
                story.append(Paragraph(get_clean_html_text(text), style_normal))
                
                # Check if this section relates to variables & operacionalizacion
                title_lower = title.lower()
                if ("variable" in title_lower or "operacionaliz" in title_lower) and not variables_rendered:
                    vars_data = c2.get("variables", {})
                    op_table_data = vars_data.get("operacionalizacion_tabla", []) if vars_data else []
                    if op_table_data:
                        story.append(Spacer(1, 0.3*cm))
                        story.append(Paragraph("<i>Tabla 1. Matriz de Operacionalización de Variables</i>", style_normal))
                        t_data = [['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']]
                        for r in op_table_data:
                            t_data.append([
                                r.get("variable", ""), r.get("definicion_conceptual", ""),
                                r.get("definicion_operacional", ""), r.get("dimensiones", ""),
                                r.get("indicadores", ""), r.get("escala_medicion", "")
                            ])
                        t = Table(t_data, colWidths=[2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0,0), (-1,0), colors.grey),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0,0), (-1,0), 4),
                            ('GRID', (0,0), (-1,-1), 1, colors.black),
                            ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                            ('FONTSIZE', (0,0), (-1,-1), 8),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 0.5*cm))
                        variables_rendered = True

        # Append variables table if it wasn't rendered
        if not variables_rendered:
            vars_data = c2.get("variables", {})
            op_table_data = vars_data.get("operacionalizacion_tabla", []) if vars_data else []
            if op_table_data:
                story.append(Paragraph("Variables y Operacionalización", style_section))
                story.append(Paragraph("<i>Tabla 1. Matriz de Operacionalización de Variables</i>", style_normal))
                t_data = [['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']]
                for r in op_table_data:
                    t_data.append([
                        r.get("variable", ""), r.get("definicion_conceptual", ""),
                        r.get("definicion_operacional", ""), r.get("dimensiones", ""),
                        r.get("indicadores", ""), r.get("escala_medicion", "")
                    ])
                t = Table(t_data, colWidths=[2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 4),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.5*cm))
    else:
        # Fallback default sequence
        sections_c2 = [
            ("2.1 Tipo de investigación", c2.get("tipo_investigacion")),
            ("2.2 Nivel de investigación", c2.get("nivel_investigacion")),
            ("2.3 Diseño de investigación", c2.get("diseno_investigacion"))
        ]
        for title, text in sections_c2:
            if text:
                story.append(Paragraph(title, style_section))
                story.append(Paragraph(text, style_normal))

        # Población, muestra y muestreo
        p_pm = c2.get("poblacion_muestra", {})
        if p_pm:
            story.append(Paragraph("2.4 Población, muestra y muestreo", style_section))
            for sub_t, key in [
                ("2.4.1 Población", "poblacion"),
                ("2.4.2 Muestra", "muestra"),
                ("2.4.3 Muestreo", "muestreo")
            ]:
                val = p_pm.get(key)
                if val:
                    story.append(Paragraph(f"<b>{sub_t}</b>", style_normal))
                    story.append(Paragraph(val, style_normal))

        # Variables y operacionalización
        vars_data = c2.get("variables", {})
        if vars_data:
            story.append(Paragraph("2.5 Variables", style_section))
            if vars_data.get("tipo"):
                story.append(Paragraph("<b>2.5.1 Tipo de Variables</b>", style_normal))
                story.append(Paragraph(vars_data.get("tipo"), style_normal))
                
            op_table_data = vars_data.get("operacionalizacion_tabla", [])
            if op_table_data:
                story.append(Paragraph("<b>2.5.2 Operacionalización de Variables</b>", style_normal))
                story.append(Spacer(1, 0.3*cm))
                story.append(Paragraph("<i>Tabla 1. Matriz de Operacionalización de Variables</i>", style_normal))
                
                t_data = [['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']]
                for r in op_table_data:
                    t_data.append([
                        r.get("variable", ""), r.get("definicion_conceptual", ""),
                        r.get("definicion_operacional", ""), r.get("dimensiones", ""),
                        r.get("indicadores", ""), r.get("escala_medicion", "")
                    ])
                
                t = Table(t_data, colWidths=[2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 4),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.5*cm))

        # Técnicas e instrumentos
        ti_data = c2.get("tecnicas_instrumentos", {})
        if ti_data:
            story.append(Paragraph("2.6 Técnicas e instrumentos", style_section))
            if ti_data.get("descripcion"):
                story.append(Paragraph("<b>2.6.1 Técnicas e instrumentos de recolección</b>", style_normal))
                story.append(Paragraph(ti_data.get("descripcion"), style_normal))
            if ti_data.get("validacion_confiabilidad"):
                story.append(Paragraph("<b>2.6.2 Validación y confiabilidad</b>", style_normal))
                story.append(Paragraph(ti_data.get("validacion_confiabilidad"), style_normal))

        sections_c2_end = [
            ("2.7 Método de análisis de datos", c2.get("metodo_analisis")),
            ("2.8 Procedimiento", c2.get("procedimiento")),
            ("2.9 Consideraciones éticas", c2.get("consideraciones_eticas"))
        ]
        for title, text in sections_c2_end:
            if text:
                story.append(Paragraph(title, style_section))
                story.append(Paragraph(text, style_normal))
    story.append(PageBreak())

    # Capítulo III
    c3 = data.get("capitulo3", {})
    story.append(Paragraph("CAPÍTULO III: ASPECTOS ADMINISTRATIVOS", style_header))
    story.append(Spacer(1, 1*cm))
    
    sections_c3 = c3.get("secciones", [])
    if sections_c3 and isinstance(sections_c3, list):
        budget_rendered = False
        cronograma_rendered = False
        for sec in sections_c3:
            title = sec.get("titulo", "")
            text = sec.get("contenido", "")
            if title and text:
                story.append(Paragraph(title, style_section))
                story.append(Paragraph(get_clean_html_text(text), style_normal))
                
                title_lower = title.lower()
                
                # Render budget table if matches title
                if ("presupuesto" in title_lower or "costo" in title_lower or "financiamiento" in title_lower) and not budget_rendered:
                    pres_table = c3.get("presupuesto_tabla", [])
                    if pres_table:
                        story.append(Spacer(1, 0.3*cm))
                        story.append(Paragraph("<i>Tabla 2. Presupuesto detallado del proyecto</i>", style_normal))
                        
                        t_data = [['Categoría', 'Recurso', 'Unidad', 'Costo Unit.', 'Cant.', 'Total']]
                        total_p = 0
                        for r in pres_table:
                            c_unit = r.get("costo_unitario", 0)
                            cant = r.get("cantidad", 0)
                            costo_tot = r.get("costo_total", 0)
                            if costo_tot == 0:
                                costo_tot = c_unit * cant
                            total_p += costo_tot
                            t_data.append([
                                r.get("categoria", ""), r.get("recurso", ""), r.get("unidad", ""),
                                f"S/. {c_unit}", str(cant), f"S/. {costo_tot}"
                            ])
                        t_data.append(["TOTAL PRESUPUESTO", "", "", "", "", f"S/. {total_p}"])
                        t = Table(t_data, colWidths=[2.5*cm, 4*cm, 2*cm, 2.5*cm, 1.5*cm, 2.5*cm])
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0,0), (-1,0), colors.grey),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                            ('GRID', (0,0), (-1,-1), 1, colors.black),
                            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                            ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                            ('FONTSIZE', (0,0), (-1,-1), 8),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 0.5*cm))
                        budget_rendered = True
                        
                # Render cronograma table if matches title
                if ("cronograma" in title_lower or "actividad" in title_lower or "gantt" in title_lower or "ejecución" in title_lower or "ejecucion" in title_lower) and not cronograma_rendered:
                    crono_data = c3.get("cronograma", {})
                    crono_table = crono_data.get("cronograma_tabla", []) if crono_data else []
                    if crono_table:
                        story.append(Spacer(1, 0.3*cm))
                        story.append(Paragraph("<i>Tabla 3. Cronograma Gantt de actividades</i>", style_normal))
                        t_data = [['Actividad', 'M1', 'M2', 'M3', 'M4', 'M5', 'M6']]
                        for r in crono_table:
                            t_data.append([
                                r.get("actividad", ""), r.get("mes_1", ""), r.get("mes_2", ""),
                                r.get("mes_3", ""), r.get("mes_4", ""), r.get("mes_5", ""),
                                r.get("mes_6", "")
                            ])
                        t = Table(t_data, colWidths=[6*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm])
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0,0), (-1,0), colors.grey),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                            ('ALIGN', (1,0), (-1,-1), 'CENTER'),
                            ('GRID', (0,0), (-1,-1), 1, colors.black),
                            ('FONTSIZE', (0,0), (-1,-1), 8),
                        ]))
                        story.append(t)
                        story.append(Spacer(1, 0.5*cm))
                        cronograma_rendered = True
                        
        if not budget_rendered:
            pres_table = c3.get("presupuesto_tabla", [])
            if pres_table:
                story.append(Paragraph("Presupuesto", style_section))
                story.append(Paragraph("<i>Tabla 2. Presupuesto detallado del proyecto</i>", style_normal))
                t_data = [['Categoría', 'Recurso', 'Unidad', 'Costo Unit.', 'Cant.', 'Total']]
                total_p = 0
                for r in pres_table:
                    c_unit = r.get("costo_unitario", 0)
                    cant = r.get("cantidad", 0)
                    costo_tot = r.get("costo_total", 0)
                    if costo_tot == 0:
                        costo_tot = c_unit * cant
                    total_p += costo_tot
                    t_data.append([
                        r.get("categoria", ""), r.get("recurso", ""), r.get("unidad", ""),
                        f"S/. {c_unit}", str(cant), f"S/. {costo_tot}"
                    ])
                t_data.append(["TOTAL PRESUPUESTO", "", "", "", "", f"S/. {total_p}"])
                t = Table(t_data, colWidths=[2.5*cm, 4*cm, 2*cm, 2.5*cm, 1.5*cm, 2.5*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.5*cm))
                
        if not cronograma_rendered:
            crono_data = c3.get("cronograma", {})
            crono_table = crono_data.get("cronograma_tabla", []) if crono_data else []
            if crono_table:
                story.append(Paragraph("Cronograma de Actividades", style_section))
                story.append(Paragraph("<i>Tabla 3. Cronograma Gantt de actividades</i>", style_normal))
                t_data = [['Actividad', 'M1', 'M2', 'M3', 'M4', 'M5', 'M6']]
                for r in crono_table:
                    t_data.append([
                        r.get("actividad", ""), r.get("mes_1", ""), r.get("mes_2", ""),
                        r.get("mes_3", ""), r.get("mes_4", ""), r.get("mes_5", ""),
                        r.get("mes_6", "")
                    ])
                t = Table(t_data, colWidths=[6*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (1,0), (-1,-1), 'CENTER'),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.5*cm))
    else:
        # Fallback default sequence
        rec_data = c3.get("recursos", {})
        if rec_data:
            story.append(Paragraph("3.1 Recursos", style_section))
            for sub_t, key in [
                ("3.1.1 Personal", "personal"),
                ("3.1.2 Bienes", "bienes"),
                ("3.1.3 Viajes", "viajes"),
                ("3.1.4 Servicios", "services"), # or servicios
                ("3.1.5 Tecnológicos", "tecnologicos")
            ]:
                val = rec_data.get(key) or rec_data.get(key.replace("services","servicios"))
                if val:
                    story.append(Paragraph(f"<b>{sub_t}</b>", style_normal))
                    story.append(Paragraph(val, style_normal))
                    
        pres_table = c3.get("presupuesto_tabla", [])
        if pres_table:
            story.append(Paragraph("3.2 Presupuesto", style_section))
            story.append(Paragraph("<i>Tabla 2. Presupuesto detallado del proyecto</i>", style_normal))
            
            t_data = [['Categoría', 'Recurso', 'Unidad', 'Costo Unit.', 'Cant.', 'Total']]
            total_p = 0
            for r in pres_table:
                c_unit = r.get("costo_unitario", 0)
                cant = r.get("cantidad", 0)
                costo_tot = r.get("costo_total", 0)
                if costo_tot == 0:
                    costo_tot = c_unit * cant
                total_p += costo_tot
                t_data.append([
                    r.get("categoria", ""), r.get("recurso", ""), r.get("unidad", ""),
                    f"S/. {c_unit}", str(cant), f"S/. {costo_tot}"
                ])
            t_data.append(["TOTAL PRESUPUESTO", "", "", "", "", f"S/. {total_p}"])
            
            t = Table(t_data, colWidths=[2.5*cm, 4*cm, 2*cm, 2.5*cm, 1.5*cm, 2.5*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTNAME', (0,-1), (-1,-1), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 8),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.5*cm))
            
        fin_text = c3.get("financiamiento")
        if fin_text:
            story.append(Paragraph("3.3 Financiamiento", style_section))
            story.append(Paragraph(fin_text, style_normal))
            
        crono_data = c3.get("cronograma", {})
        if crono_data:
            story.append(Paragraph("3.4 Cronograma de ejecución", style_section))
            if crono_data.get("periodo"):
                story.append(Paragraph(f"<b>3.4.1 Período:</b> {crono_data.get('periodo')}", style_normal))
            story.append(Paragraph("<b>3.4.2 Cronograma de Actividades</b>", style_normal))
            
            crono_table = crono_data.get("cronograma_tabla", [])
            if crono_table:
                story.append(Paragraph("<i>Tabla 3. Cronograma Gantt de actividades</i>", style_normal))
                t_data = [['Actividad', 'M1', 'M2', 'M3', 'M4', 'M5', 'M6']]
                for r in crono_table:
                    t_data.append([
                        r.get("actividad", ""), r.get("mes_1", ""), r.get("mes_2", ""),
                        r.get("mes_3", ""), r.get("mes_4", ""), r.get("mes_5", ""),
                        r.get("mes_6", "")
                    ])
                t = Table(t_data, colWidths=[6*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm, 1.5*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.grey),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                    ('ALIGN', (1,0), (-1,-1), 'CENTER'),
                    ('GRID', (0,0), (-1,-1), 1, colors.black),
                    ('FONTSIZE', (0,0), (-1,-1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 0.5*cm))
    story.append(PageBreak())

    # Referencias
    story.append(Paragraph("REFERENCIAS BIBLIOGRÁFICAS", style_header))
    story.append(Spacer(1, 1*cm))
    refs = data.get("referencias", [])
    for ref in refs:
        story.append(Paragraph(ref, style_ref))
    story.append(PageBreak())

    # Anexos
    story.append(Paragraph("ANEXOS OBLIGATORIOS", style_header))
    story.append(Spacer(1, 1*cm))
    anexos = data.get("anexos", {})
    sections_an = [
        ("Anexo 1. Matriz de operacionalización de variables", anexos.get("anexo_1")),
        ("Anexo 2. Matriz de consistencia", anexos.get("anexo_2")),
        ("Anexo 3. Diagrama de Ishikawa", anexos.get("anexo_3")),
        ("Anexo 4. Árbol de problemas", anexos.get("anexo_4")),
        ("Anexo 5. Árbol de objetivos", anexos.get("anexo_5")),
        ("Anexo 6. Instrumentos de recolección de datos", anexos.get("anexo_6")),
        ("Anexo 7. Constancia de aplicación de instrumentos", anexos.get("anexo_7")),
        ("Anexo 8. Declaración de originalidad y conformidad", anexos.get("anexo_8"))
    ]
    for title, text in sections_an:
        if text:
            story.append(Paragraph(title, style_section))
            elements = parse_markdown_to_elements(text)
            append_markdown_elements_to_pdf_story(story, elements, style_normal, style_section)

    # Build PDF
    try:
        pdf.build(story, canvasmaker=NumberedCanvas)
    finally:
        for fpath in temp_files_to_clean:
            try:
                os.unlink(fpath)
            except:
                pass


def build_txt(data, output_path):
    meta = data.get("metadata", {})
    prelims = data.get("preliminares", {})
    c1 = data.get("capitulo1", {})
    c2 = data.get("capitulo2", {})
    c3 = data.get("capitulo3", {})
    refs = data.get("referencias", [])
    anexos = data.get("anexos", {})
    
    with open(output_path, "w", encoding="utf-8") as f:
        # Cover
        template_styles = data.get("templateStyles", {}) or {}
        structure = template_styles.get("structure", {}) or {}
        institution = (structure.get("institution") or meta.get("universidad") or "UNIVERSIDAD NACIONAL DE TRUJILLO").upper()
        faculty = (structure.get("faculty") or meta.get("facultad") or "FACULTAD DE INGENIERÍA").upper()
        school = (structure.get("school") or meta.get("escuela") or "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS").upper()

        f.write(f"{institution}\n")
        f.write(f"{faculty}\n")
        f.write(f"{school}\n\n")
        f.write(f"TEMA: {meta.get('titulo_proyecto', '').upper()}\n\n")
        f.write(f"Autor: {meta.get('nombre_autor', '')}\n")
        f.write(f"Asesor: Dr. {meta.get('nombre_asesor', '')}\n")
        f.write(f"Línea de Investigación: {meta.get('linea_investigacion', '')}\n")
        f.write(f"Ciudad: {meta.get('ciudad', '')} - Año: {meta.get('anio', '')}\n")
        f.write("=" * 60 + "\n\n")
        
        # Jurado
        f.write("JURADO DICTAMINADOR\n\n")
        f.write("Presidente: Dr. Roberto Carlos Medina\n")
        f.write("Secretario: Dr. Julio César Alvarez\n")
        f.write(f"Vocal (Asesor): Dr. {meta.get('nombre_asesor', '')}\n")
        f.write("=" * 60 + "\n\n")
        
        # Preliminares
        if prelims.get("dedicatoria"):
            f.write(f"DEDICATORIA\n\n{prelims.get('dedicatoria')}\n\n")
        if prelims.get("agradecimientos"):
            f.write(f"AGRADECIMIENTOS\n\n{prelims.get('agradecimientos')}\n\n")
        if prelims.get("presentacion"):
            f.write(f"PRESENTACIÓN\n\n{prelims.get('presentacion')}\n\n")
        if prelims.get("resumen"):
            f.write(f"RESUMEN\n\n{prelims.get('resumen')}\n")
            f.write(f"Palabras clave: {prelims.get('palabras_clave', '')}\n\n")
        if prelims.get("abstract"):
            f.write(f"ABSTRACT\n\n{prelims.get('abstract')}\n")
            f.write(f"Keywords: {prelims.get('keywords', '')}\n\n")
        f.write("=" * 60 + "\n\n")
            
        # Capitulo I
        f.write("CAPÍTULO I: INTRODUCCIÓN\n\n")
        sections_c1 = c1.get("secciones", [])
        if sections_c1 and isinstance(sections_c1, list):
            for sec in sections_c1:
                f.write(f"{sec.get('titulo')}\n{get_clean_text(sec.get('contenido'))}\n\n")
        else:
            sections = [
                ("1.1 Realidad problemática", c1.get("realidad_problematica")),
                ("1.2 Antecedentes de la investigación", c1.get("antecedentes")),
                ("1.3 Marco teórico", c1.get("marco_teorico")),
                ("1.4 Metodologías alternativas", c1.get("metodologias_alternativas")),
                ("1.5 Justificación de la investigación", c1.get("justificacion")),
                ("1.6 Formulación del problema", c1.get("formulacion_problema")),
                ("1.7 Hipótesis", c1.get("hipotesis")),
                ("1.8 Objetivos", c1.get("objetivos")),
                ("1.9 Limitaciones del estudio", c1.get("limitaciones"))
            ]
            for title, val in sections:
                f.write(f"{title}\n{get_clean_text(val)}\n\n")
        f.write("=" * 60 + "\n\n")

        # Capitulo II
        f.write("CAPÍTULO II: MÉTODO\n\n")
        sections_c2 = c2.get("secciones", [])
        if sections_c2 and isinstance(sections_c2, list):
            variables_rendered = False
            for sec in sections_c2:
                title = sec.get("titulo", "")
                text = sec.get("contenido", "")
                f.write(f"{title}\n{get_clean_text(text)}\n\n")
                
                title_lower = title.lower()
                if ("variable" in title_lower or "operacionaliz" in title_lower) and not variables_rendered:
                    vars_data = c2.get("variables", {})
                    f.write("2.5.2 Operacionalización de Variables:\n")
                    for row in vars_data.get("operacionalizacion_tabla", []):
                        f.write(f"- Variable: {row.get('variable')}\n")
                        f.write(f"  Def. Conceptual: {row.get('definicion_conceptual')}\n")
                        f.write(f"  Def. Operacional: {row.get('definicion_operacional')}\n")
                        f.write(f"  Dimensiones: {row.get('dimensiones')}\n")
                        f.write(f"  Indicadores: {row.get('indicadores')}\n")
                        f.write(f"  Escala: {row.get('escala_medicion')}\n")
                    f.write("\n")
                    variables_rendered = True
            if not variables_rendered:
                vars_data = c2.get("variables", {})
                f.write("2.5.2 Operacionalización de Variables:\n")
                for row in vars_data.get("operacionalizacion_tabla", []):
                    f.write(f"- Variable: {row.get('variable')}\n")
                    f.write(f"  Def. Conceptual: {row.get('definicion_conceptual')}\n")
                    f.write(f"  Def. Operacional: {row.get('definicion_operacional')}\n")
                    f.write(f"  Dimensiones: {row.get('dimensiones')}\n")
                    f.write(f"  Indicadores: {row.get('indicadores')}\n")
                    f.write(f"  Escala: {row.get('escala_medicion')}\n")
                f.write("\n")
        else:
            f.write(f"2.1 Tipo de investigación\n{c2.get('tipo_investigacion','')}\n\n")
            f.write(f"2.2 Nivel de investigación\n{c2.get('nivel_investigacion','')}\n\n")
            f.write(f"2.3 Diseño de investigación\n{c2.get('diseno_investigacion','')}\n\n")
            
            p_pm = c2.get("poblacion_muestra", {})
            f.write(f"2.4 Población, muestra y muestreo\n")
            f.write(f"2.4.1 Población: {p_pm.get('poblacion','')}\n")
            f.write(f"2.4.2 Muestra: {p_pm.get('muestra','')}\n")
            f.write(f"2.4.3 Muestreo: {p_pm.get('muestreo','')}\n\n")
            
            vars_data = c2.get("variables", {})
            f.write(f"2.5 Variables\n")
            f.write(f"2.5.1 Tipo: {vars_data.get('tipo','')}\n\n")
            f.write("2.5.2 Operacionalización de Variables:\n")
            for row in vars_data.get("operacionalizacion_tabla", []):
                f.write(f"- Variable: {row.get('variable')}\n")
                f.write(f"  Def. Conceptual: {row.get('definicion_conceptual')}\n")
                f.write(f"  Def. Operacional: {row.get('definicion_operacional')}\n")
                f.write(f"  Dimensiones: {row.get('dimensiones')}\n")
                f.write(f"  Indicadores: {row.get('indicadores')}\n")
                f.write(f"  Escala: {row.get('escala_medicion')}\n")
            f.write("\n")
            
            ti_data = c2.get("tecnicas_instrumentos", {})
            f.write("2.6 Técnicas e instrumentos\n")
            f.write(f"2.6.1 Técnicas e instrumentos: {ti_data.get('descripcion','')}\n")
            f.write(f"2.6.2 Validación y confiabilidad: {ti_data.get('validacion_confiabilidad','')}\n\n")
            
            f.write(f"2.7 Método de análisis de datos\n{c2.get('metodo_analisis','')}\n\n")
            f.write(f"2.8 Procedimiento\n{c2.get('procedimiento','')}\n\n")
            f.write(f"2.9 Consideraciones éticas\n{c2.get('consideraciones_eticas','')}\n\n")
        f.write("=" * 60 + "\n\n")

        # Capitulo III
        f.write("CAPÍTULO III: ASPECTOS ADMINISTRATIVOS\n\n")
        sections_c3 = c3.get("secciones", [])
        if sections_c3 and isinstance(sections_c3, list):
            budget_rendered = False
            cronograma_rendered = False
            for sec in sections_c3:
                title = sec.get("titulo", "")
                text = sec.get("contenido", "")
                f.write(f"{title}\n{get_clean_text(text)}\n\n")
                
                title_lower = title.lower()
                if ("presupuesto" in title_lower or "costo" in title_lower or "financiamiento" in title_lower) and not budget_rendered:
                    f.write("3.2 Presupuesto:\n")
                    for r in c3.get("presupuesto_tabla", []):
                        f.write(f"- [{r.get('categoria')}] {r.get('recurso')}: {r.get('unidad')} S/.{r.get('costo_unitario')} x {r.get('cantidad')} = S/.{r.get('costo_total')}\n")
                    f.write("\n")
                    budget_rendered = True
                    
                if ("cronograma" in title_lower or "actividad" in title_lower or "gantt" in title_lower or "ejecución" in title_lower or "ejecucion" in title_lower) and not cronograma_rendered:
                    crono_data = c3.get("cronograma", {})
                    f.write("Actividades:\n")
                    for row in crono_data.get("cronograma_tabla", []):
                        f.write(f"- {row.get('actividad')}: Meses [{row.get('mes_1','')}{row.get('mes_2','')}{row.get('mes_3','')}{row.get('mes_4','')}{row.get('mes_5','')}{row.get('mes_6','')}]\n")
                    f.write("\n")
                    cronograma_rendered = True
                    
            if not budget_rendered:
                f.write("3.2 Presupuesto:\n")
                for r in c3.get("presupuesto_tabla", []):
                    f.write(f"- [{r.get('categoria')}] {r.get('recurso')}: {r.get('unidad')} S/.{r.get('costo_unitario')} x {r.get('cantidad')} = S/.{r.get('costo_total')}\n")
                f.write("\n")
            if not cronograma_rendered:
                crono_data = c3.get("cronograma", {})
                f.write("Actividades:\n")
                for row in crono_data.get("cronograma_tabla", []):
                    f.write(f"- {row.get('actividad')}: Meses [{row.get('mes_1','')}{row.get('mes_2','')}{row.get('mes_3','')}{row.get('mes_4','')}{row.get('mes_5','')}{row.get('mes_6','')}]\n")
                f.write("\n")
        else:
            rec_data = c3.get("recursos", {})
            f.write("3.1 Recursos\n")
            f.write(f"3.1.1 Personal: {rec_data.get('personal','')}\n")
            f.write(f"3.1.2 Bienes: {rec_data.get('bienes','')}\n")
            f.write(f"3.1.3 Viajes: {rec_data.get('viajes','')}\n")
            f.write(f"3.1.4 Servicios: {rec_data.get('servicios','')}\n")
            f.write(f"3.1.5 Tecnológicos: {rec_data.get('tecnologicos','')}\n\n")
            
            f.write("3.2 Presupuesto:\n")
            for r in c3.get("presupuesto_tabla", []):
                f.write(f"- [{r.get('categoria')}] {r.get('recurso')}: {r.get('unidad')} S/.{r.get('costo_unitario')} x {r.get('cantidad')} = S/.{r.get('costo_total')}\n")
            f.write("\n")
            
            f.write(f"3.3 Financiamiento: {c3.get('financiamiento','')}\n\n")
            
            crono_data = c3.get("cronograma", {})
            f.write(f"3.4 Cronograma de ejecución\n")
            f.write(f"3.4.1 Período: {crono_data.get('periodo','')}\n\n")
            f.write("Actividades:\n")
            for row in crono_data.get("cronograma_tabla", []):
                f.write(f"- {row.get('actividad')}: Meses [{row.get('mes_1','')}{row.get('mes_2','')}{row.get('mes_3','')}{row.get('mes_4','')}{row.get('mes_5','')}{row.get('mes_6','')}]\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # Referencias
        f.write("REFERENCIAS BIBLIOGRÁFICAS\n\n")
        for ref in refs:
            f.write(f"{ref}\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # Anexos
        f.write("ANEXOS OBLIGATORIOS\n\n")
        for key, val in [
            ("Anexo 1. Matriz de operacionalización de variables", anexos.get("anexo_1")),
            ("Anexo 2. Matriz de consistencia", anexos.get("anexo_2")),
            ("Anexo 3. Diagrama de Ishikawa", anexos.get("anexo_3")),
            ("Anexo 4. Árbol de problemas", anexos.get("anexo_4")),
            ("Anexo 5. Árbol de objetivos", anexos.get("anexo_5")),
            ("Anexo 6. Instrumentos de recolección de datos", anexos.get("anexo_6")),
            ("Anexo 7. Constancia de aplicación de instrumentos", anexos.get("anexo_7")),
            ("Anexo 8. Declaración de originalidad y conformidad", anexos.get("anexo_8"))
        ]:
            f.write(f"{key}\n{val}\n\n")

# ---------------------------------------------------------------------------
# Final Thesis Export Functions (5-Chapter Structure with past tense)
# ---------------------------------------------------------------------------

def build_final_thesis_docx(data, output_path):
    doc = Document()
    
    # ------------------ COVER PAGE (SECTION 1) ------------------
    template_styles = data.get("templateStyles", {}) or {}
    section1 = doc.sections[0]
    apply_section_margins(section1, template_styles, default_left=3.0)
    
    meta = data.get("metadata", {})
    
    structure = template_styles.get("structure", {}) or {}
    institution = (structure.get("institution") or meta.get("universidad") or "UNIVERSIDAD NACIONAL DE TRUJILLO").upper()
    faculty = (structure.get("faculty") or meta.get("facultad") or "FACULTAD DE INGENIERÍA").upper()
    school = (structure.get("school") or meta.get("escuela") or "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS").upper()
    
    school_clean = school.replace("ESCUELA PROFESIONAL DE ", "").replace("ESCUELA DE ", "").replace("CARRERA DE ", "").replace("PROGRAMA DE ", "")
    if school_clean.startswith("INGENIERÍA DE "):
        optar_titulo = "INGENIERO DE " + school_clean[14:]
    elif school_clean.startswith("INGENIERÍA "):
        optar_titulo = "INGENIERO " + school_clean[11:]
    else:
        optar_titulo = school_clean

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{institution}\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    run2 = p.add_run(f"{faculty}\n{school}\n\n")
    apply_text_formatting(run2, font_name="Arial Narrow", size=14, bold=True)
    
    logo_data = template_styles.get("logo")
    logo_inserted = False
    if logo_data and logo_data.get("base64"):
        try:
            logo_bytes = base64.b64decode(logo_data["base64"])
            ext = logo_data.get("ext", "png")
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp_img:
                tmp_img.write(logo_bytes)
                tmp_img_path = tmp_img.name
            
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(tmp_img_path, width=Inches(1.8))
            os.unlink(tmp_img_path)
            logo_inserted = True
        except Exception as img_err:
            print(f"Error inserting logo in final thesis docx: {img_err}", file=sys.stderr)
            
    p_spacing = doc.add_paragraph()
    p_spacing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing_text = "\n\n" if logo_inserted else "\n\n\n\n"
    run_space = p_spacing.add_run(spacing_text)
    apply_text_formatting(run_space, font_name="Arial Narrow", size=12)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"\"{meta.get('titulo_proyecto', '').upper()}\"\n\n")
    apply_text_formatting(run_title, font_name="Arial Narrow", size=14, bold=True)
    
    p_spacing2 = doc.add_paragraph()
    p_spacing2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing_text2 = "\n" if logo_inserted else "\n\n\n"
    run_space2 = p_spacing2.add_run(spacing_text2)
    apply_text_formatting(run_space2, font_name="Arial Narrow", size=12)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run(f"TESIS PARA OPTAR EL TÍTULO PROFESIONAL DE\n{optar_titulo}\n\n\n")
    apply_text_formatting(run_sub, font_name="Arial Narrow", size=12, bold=True)
    
    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_author.add_run(f"Autor: {meta.get('nombre_autor', '')}\nAsesor: Dr. {meta.get('nombre_asesor', '')}\n\n")
    apply_text_formatting(run_auth, font_name="Arial Narrow", size=12)
    
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run(f"Línea de Investigación: {meta.get('linea_investigacion', '')}\n\n\n\n")
    apply_text_formatting(run_line, font_name="Arial Narrow", size=11, italic=True)
    
    p_city = doc.add_paragraph()
    p_city.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_city = p_city.add_run(f"{meta.get('ciudad', '')} - Perú\n{meta.get('anio', '')}")
    apply_text_formatting(run_city, font_name="Arial Narrow", size=12, bold=True)
    
    # ------------------ PRELIMINARES (SECTION 2) ------------------
    section2 = doc.add_section()
    apply_section_margins(section2, template_styles, default_left=3.0)
    section2.footer.is_linked_to_previous = False
    configure_page_numbering(section2, fmt="romanLower", start=2)
    p_footer2 = section2.footer.paragraphs[0]
    p_footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number_to_footer(p_footer2)

    # Jurado Dictaminador
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\nJURADO DICTAMINADOR\n\n\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    
    jurados = [
        ("Presidente", "Dr. Roberto Carlos Medina"),
        ("Secretario", "Dr. Julio César Alvarez"),
        ("Vocal (Asesor)", f"Dr. {meta.get('nombre_asesor', '')}")
    ]
    for role, name in jurados:
        p_j = doc.add_paragraph()
        p_j.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_j.paragraph_format.space_before = Pt(30)
        run_line = p_j.add_run("_______________________________\n")
        apply_text_formatting(run_line, font_name="Arial Narrow", size=12)
        run_name = p_j.add_run(f"{name}\n")
        apply_text_formatting(run_name, font_name="Arial Narrow", size=12, bold=True)
        run_role = p_j.add_run(role)
        apply_text_formatting(run_role, font_name="Arial Narrow", size=11, italic=True)
        
    doc.add_page_break()

    # Dedicatoria & Agradecimientos
    prelims = data.get("preliminares", {})
    if prelims.get("dedicatoria"):
        p = doc.add_paragraph()
        run = p.add_run("DEDICATORIA\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_ded = doc.add_paragraph()
        p_ded.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_ded.paragraph_format.left_indent = Cm(5.0)
        run_ded = p_ded.add_run(prelims.get("dedicatoria"))
        apply_text_formatting(run_ded, font_name="Arial Narrow", italic=True)
        set_paragraph_spacing(p_ded)
        doc.add_page_break()
        
    if prelims.get("agradecimientos"):
        p = doc.add_paragraph()
        run = p.add_run("AGRADECIMIENTOS\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_agr = doc.add_paragraph()
        run_agr = p_agr.add_run(prelims.get("agradecimientos"))
        apply_text_formatting(run_agr, font_name="Arial Narrow")
        set_paragraph_spacing(p_agr)
        doc.add_page_break()

    # Índice General
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE GENERAL\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    
    for item in prelims.get("indice_general", []):
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.tab_stops.add_tab_stop(Cm(14.0), alignment=2)
        run_item = p_toc.add_run(f"{item}")
        apply_text_formatting(run_item, font_name="Arial Narrow")
        set_paragraph_spacing(p_toc, line_spacing=1.2, before=0, after=2)
        
    doc.add_page_break()

    # Índice de tablas
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE DE TABLAS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    for idx, item in enumerate(prelims.get("indice_tablas", [])):
        p_tbl = doc.add_paragraph()
        run_tbl = p_tbl.add_run(item)
        apply_text_formatting(run_tbl, font_name="Arial Narrow")
        set_paragraph_spacing(p_tbl, line_spacing=1.2, before=0, after=2)
    doc.add_page_break()

    # Índice de figuras
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE DE FIGURAS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
    for idx, item in enumerate(prelims.get("indice_figuras", prelims.get("indice_figures", []))):
        p_fig = doc.add_paragraph()
        run_fig = p_fig.add_run(item)
        apply_text_formatting(run_fig, font_name="Arial Narrow")
        set_paragraph_spacing(p_fig, line_spacing=1.2, before=0, after=2)
    doc.add_page_break()

    # Presentación
    if prelims.get("presentacion"):
        p = doc.add_paragraph()
        run = p.add_run("PRESENTACIÓN\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_pres = doc.add_paragraph()
        run_pres = p_pres.add_run(prelims.get("presentacion"))
        apply_text_formatting(run_pres, font_name="Arial Narrow")
        set_paragraph_spacing(p_pres)
        doc.add_page_break()
        
    # Resumen
    if prelims.get("resumen"):
        p = doc.add_paragraph()
        run = p.add_run("RESUMEN\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_res = doc.add_paragraph()
        run_res = p_res.add_run(prelims.get("resumen"))
        apply_text_formatting(run_res, font_name="Arial Narrow")
        set_paragraph_spacing(p_res)
        
        if prelims.get("palabras_clave"):
            p_kw = doc.add_paragraph()
            run_lbl = p_kw.add_run("Palabras clave: ")
            apply_text_formatting(run_lbl, font_name="Arial Narrow", bold=True)
            run_val = p_kw.add_run(prelims.get("palabras_clave"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_kw, before=12)
            
        doc.add_page_break()
        
    # Abstract
    if prelims.get("abstract"):
        p = doc.add_paragraph()
        run = p.add_run("ABSTRACT\n\n")
        apply_text_formatting(run, font_name="Arial Narrow", size=14, bold=True)
        p_abs = doc.add_paragraph()
        run_abs = p_abs.add_run(prelims.get("abstract"))
        apply_text_formatting(run_abs, font_name="Arial Narrow")
        set_paragraph_spacing(p_abs)
        
        if prelims.get("keywords"):
            p_kw = doc.add_paragraph()
            run_lbl = p_kw.add_run("Keywords: ")
            apply_text_formatting(run_lbl, font_name="Arial Narrow", bold=True)
            run_val = p_kw.add_run(prelims.get("keywords"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_kw, before=12)
            
    # ------------------ MAIN BODY (SECTION 3) ------------------
    section3 = doc.add_section()
    apply_section_margins(section3, template_styles, default_left=3.0)
    section3.footer.is_linked_to_previous = False
    configure_page_numbering(section3, fmt="decimal", start=1)
    p_footer3 = section3.footer.paragraphs[0]
    p_footer3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number_to_footer(p_footer3)

    # CAPÍTULO I: INTRODUCCIÓN
    c1 = data.get("capitulo1", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO I: INTRODUCCIÓN\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    sections_c1 = [
        ("1.1 Realidad problemática", c1.get("realidad_problematica")),
        ("1.2 Antecedentes de la investigación", c1.get("antecedentes")),
        ("1.3 Justificación de la investigación", c1.get("justificacion")),
        ("1.4 Formulación del problema", c1.get("formulacion_problema")),
        ("1.5 Hipótesis", c1.get("hipotesis")),
        ("1.6 Objetivos", c1.get("objetivos")),
        ("1.7 Limitaciones del estudio", c1.get("limitaciones"))
    ]
    for title, text in sections_c1:
        if text:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run(title)
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            
            p_text = doc.add_paragraph()
            run_txt = p_text.add_run(get_clean_text(text))
            apply_text_formatting(run_txt, font_name="Arial Narrow")
            set_paragraph_spacing(p_text)
                
    doc.add_page_break()

    # CAPÍTULO II: MARCO TEÓRICO
    c2 = data.get("capitulo2", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO II: MARCO TEÓRICO\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    if c2.get("bases_teoricas"):
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("2.1 Bases teóricas")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        p_text = doc.add_paragraph()
        run_txt = p_text.add_run(c2.get("bases_teoricas"))
        apply_text_formatting(run_txt, font_name="Arial Narrow")
        set_paragraph_spacing(p_text)
        
    doc.add_page_break()

    # CAPÍTULO III: MÉTODO
    c3 = data.get("capitulo3", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO III: MÉTODO\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    sections_c3 = [
        ("3.1 Tipo de investigación", c3.get("tipo_investigacion")),
        ("3.2 Nivel de investigación", c3.get("nivel_investigacion")),
        ("3.3 Diseño de investigación", c3.get("diseno_investigacion"))
    ]
    for title, text in sections_c3:
        if text:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run(title)
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            p_text = doc.add_paragraph()
            run_txt = p_text.add_run(text)
            apply_text_formatting(run_txt, font_name="Arial Narrow")
            set_paragraph_spacing(p_text)

    # Población y muestra
    p_pm = c3.get("poblacion_muestra", {})
    if p_pm:
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("3.4 Población, muestra y muestreo")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        for label, key in [
            ("3.4.1 Población", "poblacion"),
            ("3.4.2 Muestra", "muestra"),
            ("3.4.3 Muestreo", "muestreo")
        ]:
            val = p_pm.get(key)
            if val:
                p_sub = doc.add_paragraph()
                run_sub = p_sub.add_run(label)
                apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
                set_paragraph_spacing(p_sub, before=8, after=4)
                
                p_val = doc.add_paragraph()
                run_val = p_val.add_run(val)
                apply_text_formatting(run_val, font_name="Arial Narrow")
                set_paragraph_spacing(p_val)

    # Variables y matriz
    vars_data = c3.get("variables", {})
    if vars_data:
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("3.5 Variables")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        if vars_data.get("tipo"):
            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run("3.5.1 Tipo de Variables")
            apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_sub, before=8, after=4)
            p_val = doc.add_paragraph()
            run_val = p_val.add_run(vars_data.get("tipo"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_val)
            
        op_table_data = vars_data.get("operacionalizacion_tabla", [])
        if op_table_data:
            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run("3.5.2 Operacionalización de Variables")
            apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_sub, before=8, after=4)
            
            p_lbl = doc.add_paragraph()
            run_lbl = p_lbl.add_run("Tabla 1. Matriz de Operacionalización de Variables")
            apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
            
            table = doc.add_table(rows=1, cols=6)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            hdr_titles = ['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']
            for idx, col_title in enumerate(hdr_titles):
                hdr_cells[idx].text = col_title
                apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            for r in op_table_data:
                row_cells = table.add_row().cells
                row_cells[0].text = r.get("variable", "")
                row_cells[1].text = r.get("definicion_conceptual", "")
                row_cells[2].text = r.get("definicion_operacional", "")
                row_cells[3].text = r.get("dimensiones", "")
                row_cells[4].text = r.get("indicadores", "")
                row_cells[5].text = r.get("escala_medicion", "")
                for cell in row_cells:
                    apply_text_formatting(cell.paragraphs[0].runs[0], font_name="Arial Narrow", size=9)

    # Técnicas e instrumentos
    ti_data = c3.get("tecnicas_instrumentos", {})
    if ti_data:
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("3.6 Técnicas e instrumentos de recolección de datos")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        if ti_data.get("descripcion"):
            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run("3.6.1 Descripción de técnicas e instrumentos")
            apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_sub, before=8, after=4)
            p_val = doc.add_paragraph()
            run_val = p_val.add_run(ti_data.get("descripcion"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_val)
            
        if ti_data.get("validacion_confiabilidad"):
            p_sub = doc.add_paragraph()
            run_sub = p_sub.add_run("3.6.2 Validación y confiabilidad")
            apply_text_formatting(run_sub, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_sub, before=8, after=4)
            p_val = doc.add_paragraph()
            run_val = p_val.add_run(ti_data.get("validacion_confiabilidad"))
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_val)

    sections_c3_end = [
        ("3.7 Método de análisis de datos", c3.get("metodo_analisis")),
        ("3.8 Procedimiento realizado", c3.get("procedimiento")),
        ("3.9 Consideraciones éticas", c3.get("consideraciones_eticas"))
    ]
    for title, text in sections_c3_end:
        if text:
            p_title = doc.add_paragraph()
            run_t = p_title.add_run(title)
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_title, before=12, after=6)
            p_text = doc.add_paragraph()
            run_txt = p_text.add_run(text)
            apply_text_formatting(run_txt, font_name="Arial Narrow")
            set_paragraph_spacing(p_text)

    doc.add_page_break()

    # CAPÍTULO IV: RESULTADOS
    c4 = data.get("capitulo4", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO IV: RESULTADOS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)

    if c4.get("resultados"):
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("4.1 Análisis descriptivo y contrastación de hipótesis")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        p_text = doc.add_paragraph()
        run_txt = p_text.add_run(c4.get("resultados"))
        apply_text_formatting(run_txt, font_name="Arial Narrow")
        set_paragraph_spacing(p_text)

    # Render results tables
    if c4.get("resultados_tablas"):
        for idx, tab in enumerate(c4.get("resultados_tablas")):
            p_tbl_lbl = doc.add_paragraph()
            run_lbl = p_tbl_lbl.add_run(tab.get("titulo", f"Tabla {idx+2}"))
            apply_text_formatting(run_lbl, font_name="Arial Narrow", size=11, bold=True)
            set_paragraph_spacing(p_tbl_lbl, before=12, after=6)
            
            cols = tab.get("columnas", [])
            filas = tab.get("filas", [])
            
            table = doc.add_table(rows=1, cols=len(cols))
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table.rows[0].cells
            for col_idx, col_title in enumerate(cols):
                hdr_cells[col_idx].text = col_title
                apply_text_formatting(hdr_cells[col_idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=10, bold=True)
                hdr_cells[col_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            for r in filas:
                row_cells = table.add_row().cells
                for col_idx, col_title in enumerate(cols):
                    row_cells[col_idx].text = str(r.get(col_title, ""))
                    apply_text_formatting(row_cells[col_idx].paragraphs[0].runs[0], font_name="Arial Narrow", size=9)
            apply_three_line_table_style(table)

    # Render results figures description
    if c4.get("resultados_figuras"):
        for fig in c4.get("resultados_figuras"):
            p_fig = doc.add_paragraph()
            p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_f = p_fig.add_run(f"\n[ {fig.get('titulo', 'Figura')} ]\n")
            apply_text_formatting(run_f, font_name="Arial Narrow", size=11, bold=True)
            
            p_fig_cap = doc.add_paragraph()
            p_fig_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_fc = p_fig_cap.add_run(f"{fig.get('titulo', 'Figura')}. {fig.get('descripcion', '')}")
            apply_text_formatting(run_fc, font_name="Arial Narrow", size=10, italic=True)
            set_paragraph_spacing(p_fig_cap, before=4, after=12)

    doc.add_page_break()

    # CAPÍTULO V: DISCUSIÓN Y RECOMENDACIONES
    c5 = data.get("capitulo5", {})
    p = doc.add_paragraph()
    run = p.add_run("CAPÍTULO V: DISCUSIÓN Y RECOMENDACIONES\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)

    if c5.get("discusion"):
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("5.1 Discusión de los hallazgos")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        p_text = doc.add_paragraph()
        run_txt = p_text.add_run(c5.get("discusion"))
        apply_text_formatting(run_txt, font_name="Arial Narrow")
        set_paragraph_spacing(p_text)

    # Recommendations
    recs = c5.get("recomendaciones", [])
    if recs:
        p_title = doc.add_paragraph()
        run_t = p_title.add_run("5.2 Recomendaciones")
        apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
        set_paragraph_spacing(p_title, before=12, after=6)
        
        for idx, rec in enumerate(recs):
            p_rec = doc.add_paragraph()
            run_n = p_rec.add_run(f"{idx+1}. ")
            apply_text_formatting(run_n, font_name="Arial Narrow", bold=True)
            run_val = p_rec.add_run(rec)
            apply_text_formatting(run_val, font_name="Arial Narrow")
            set_paragraph_spacing(p_rec, before=0, after=4)

    doc.add_page_break()

    # CONCLUSIONES
    p = doc.add_paragraph()
    run = p.add_run("CONCLUSIONES\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    concs = data.get("conclusiones", [])
    for idx, conc in enumerate(concs):
        p_c = doc.add_paragraph()
        run_n = p_c.add_run(f"{idx+1}. ")
        apply_text_formatting(run_n, font_name="Arial Narrow", bold=True)
        run_val = p_c.add_run(conc)
        apply_text_formatting(run_val, font_name="Arial Narrow")
        set_paragraph_spacing(p_c, before=0, after=6)
        
    doc.add_page_break()

    # REFERENCIAS BIBLIOGRÁFICAS
    p = doc.add_paragraph()
    run = p.add_run("REFERENCIAS BIBLIOGRÁFICAS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    refs = data.get("referencias", [])
    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.left_indent = Cm(1.27)
        p_ref.paragraph_format.first_line_indent = Cm(-1.27)
        run_ref = p_ref.add_run(ref)
        apply_text_formatting(run_ref, font_name="Arial Narrow", size=11)
        set_paragraph_spacing(p_ref, line_spacing=1.5, before=0, after=6)
        
    doc.add_page_break()

    # ANEXOS OBLIGATORIOS
    p = doc.add_paragraph()
    run = p.add_run("ANEXOS OBLIGATORIOS\n\n")
    apply_text_formatting(run, font_name="Arial Narrow", size=16, bold=True)
    
    anexos = data.get("anexos", {})
    sections_an = [
        ("Anexo 1. Matriz de operacionalización de variables", anexos.get("anexo_1")),
        ("Anexo 2. Matriz de consistencia", anexos.get("anexo_2")),
        ("Anexo 3. Diagrama de Ishikawa", anexos.get("anexo_3")),
        ("Anexo 4. Árbol de problemas", anexos.get("anexo_4")),
        ("Anexo 5. Árbol de objetivos", anexos.get("anexo_5")),
        ("Anexo 6. Instrumentos de recolección de datos", anexos.get("anexo_6")),
        ("Anexo 7. Constancia de aplicación de instrumentos", anexos.get("anexo_7")),
        ("Anexo 8. Declaración de originalidad y conformidad", anexos.get("anexo_8"))
    ]
    for title, text in sections_an:
        if text:
            p_t = doc.add_paragraph()
            run_t = p_t.add_run(title)
            apply_text_formatting(run_t, font_name="Arial Narrow", size=12, bold=True)
            set_paragraph_spacing(p_t, before=12, after=6)
            
            elements = parse_markdown_to_elements(text)
            append_markdown_elements_to_docx(doc, elements)

    doc.save(output_path)


def build_final_thesis_pdf(data, output_path):
    template_styles = data.get("templateStyles", {}) or {}
    temp_files_to_clean = []
    margins = template_styles.get("margins", {})
    
    pdf = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=margins.get("left", 3.0)*cm,
        rightMargin=margins.get("right", 2.5)*cm,
        topMargin=margins.get("top", 2.5)*cm,
        bottomMargin=margins.get("bottom", 2.5)*cm
    )
    
    styles = getSampleStyleSheet()
    font_reg, font_bold, font_italic = get_reportlab_fonts(is_article=False)
    line_sp = get_line_spacing(is_article=False)
    
    style_normal = ParagraphStyle(
        'CustomNormalFinal', parent=styles['Normal'], fontName=font_reg,
        fontSize=11, leading=int(11 * line_sp * 1.2), alignment=4, spaceAfter=6
    )
    style_title = ParagraphStyle(
        'CustomTitleFinal', parent=styles['Normal'], fontName=font_bold,
        fontSize=14, leading=20, alignment=1, spaceAfter=15
    )
    style_header = ParagraphStyle(
        'CustomHeaderFinal', parent=styles['Normal'], fontName=font_bold,
        fontSize=15, leading=22, alignment=1, spaceAfter=20
    )
    style_section = ParagraphStyle(
        'CustomSectionFinal', parent=styles['Normal'], fontName=font_bold,
        fontSize=12, leading=18, alignment=0, spaceBefore=12, spaceAfter=6
    )
    style_ref = ParagraphStyle(
        'CustomRefFinal', parent=styles['Normal'], fontName=font_reg,
        fontSize=10, leading=15, alignment=4, leftIndent=1.27*cm, firstLineIndent=-1.27*cm, spaceAfter=6
    )

    story = []
    meta = data.get("metadata", {})
    
    # Cover Page
    structure = template_styles.get("structure", {}) or {}
    institution = (structure.get("institution") or meta.get("universidad") or "UNIVERSIDAD NACIONAL DE TRUJILLO").upper()
    faculty = (structure.get("faculty") or meta.get("facultad") or "FACULTAD DE INGENIERÍA").upper()
    school = (structure.get("school") or meta.get("escuela") or "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS").upper()
    
    school_clean = school.replace("ESCUELA PROFESIONAL DE ", "").replace("ESCUELA DE ", "").replace("CARRERA DE ", "").replace("PROGRAMA DE ", "")
    if school_clean.startswith("INGENIERÍA DE "):
        optar_titulo = "INGENIERO DE " + school_clean[14:]
    elif school_clean.startswith("INGENIERÍA "):
        optar_titulo = "INGENIERO " + school_clean[11:]
    else:
        optar_titulo = school_clean

    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(institution, style_title))
    story.append(Paragraph(f"{faculty}<br/>{school}", style_title))
    story.append(Spacer(1, 0.5*cm))
    
    logo_data = template_styles.get("logo")
    logo_inserted = False
    if logo_data and logo_data.get("base64"):
        try:
            logo_bytes = base64.b64decode(logo_data["base64"])
            ext = logo_data.get("ext", "png")
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tmp_img:
                tmp_img.write(logo_bytes)
                tmp_img_path = tmp_img.name
            
            temp_files_to_clean.append(tmp_img_path)
            from reportlab.platypus import Image
            logo_w = 1.8 * inch
            story.append(Image(tmp_img_path, width=logo_w, height=logo_w))
            logo_inserted = True
        except Exception as img_err:
            print(f"Error inserting logo in PDF: {img_err}", file=sys.stderr)
            
    if logo_inserted:
        story.append(Spacer(1, 0.8*cm))
    else:
        story.append(Spacer(1, 1.5*cm))
        
    story.append(Paragraph(f"\"{meta.get('titulo_proyecto', '').upper()}\"", style_header))
    
    if logo_inserted:
        story.append(Spacer(1, 0.8*cm))
    else:
        story.append(Spacer(1, 1.5*cm))
        
    story.append(Paragraph(f"TESIS PARA OPTAR EL TÍTULO PROFESIONAL DE<br/>{optar_titulo}", style_title))
    
    if logo_inserted:
        story.append(Spacer(1, 0.8*cm))
    else:
        story.append(Spacer(1, 1.5*cm))
    story.append(Paragraph(f"Autor: {meta.get('nombre_autor', '')}<br/>Asesor: Dr. {meta.get('nombre_asesor', '')}", style_normal))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(f"Línea de Investigación: {meta.get('linea_investigacion', '')}", style_normal))
    story.append(Spacer(1, 2*cm))
    story.append(Paragraph(f"{meta.get('ciudad', '')} - Perú<br/>{meta.get('anio', '')}", style_title))
    story.append(PageBreak())

    # Jurado
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("JURADO DICTAMINADOR", style_title))
    story.append(Spacer(1, 1.5*cm))
    
    jurados = [
        ("Presidente", "Dr. Roberto Carlos Medina"),
        ("Secretario", "Dr. Julio César Alvarez"),
        ("Vocal (Asesor)", f"Dr. {meta.get('nombre_asesor', '')}")
    ]
    for role, name in jurados:
        story.append(Spacer(1, 1.5*cm))
        story.append(Paragraph("_______________________________", style_title))
        story.append(Paragraph(f"<b>{name}</b><br/><i>{role}</i>", style_title))
    story.append(PageBreak())

    # Dedicatoria & Agradecimientos
    prelims = data.get("preliminares", {})
    if prelims.get("dedicatoria"):
        story.append(Paragraph("DEDICATORIA", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(f"<i>{prelims.get('dedicatoria')}</i>", style_normal))
        story.append(PageBreak())
        
    if prelims.get("agradecimientos"):
        story.append(Paragraph("AGRADECIMIENTOS", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("agradecimientos"), style_normal))
        story.append(PageBreak())

    # TOC
    story.append(Paragraph("ÍNDICE GENERAL", style_title))
    story.append(Spacer(1, 0.5*cm))
    for item in prelims.get("indice_general", []):
        story.append(Paragraph(f"{item}", style_normal))
    story.append(PageBreak())

    # Índice de tablas
    story.append(Paragraph("ÍNDICE DE TABLAS", style_title))
    story.append(Spacer(1, 0.5*cm))
    for item in prelims.get("indice_tablas", []):
        story.append(Paragraph(item, style_normal))
    story.append(PageBreak())

    # Índice de figuras
    story.append(Paragraph("ÍNDICE DE FIGURAS", style_title))
    story.append(Spacer(1, 0.5*cm))
    for item in prelims.get("indice_figuras", prelims.get("indice_figures", [])):
        story.append(Paragraph(item, style_normal))
    story.append(PageBreak())

    # Presentación
    if prelims.get("presentacion"):
        story.append(Paragraph("PRESENTACIÓN", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("presentacion"), style_normal))
        story.append(PageBreak())
        
    # Resumen y Abstract
    if prelims.get("resumen"):
        story.append(Paragraph("RESUMEN", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("resumen"), style_normal))
        if prelims.get("palabras_clave"):
            story.append(Paragraph(f"<b>Palabras clave:</b> {prelims.get('palabras_clave')}", style_normal))
        story.append(PageBreak())
        
    if prelims.get("abstract"):
        story.append(Paragraph("ABSTRACT", style_section))
        story.append(Spacer(1, 0.5*cm))
        story.append(Paragraph(prelims.get("abstract"), style_normal))
        if prelims.get("keywords"):
            story.append(Paragraph(f"<b>Keywords:</b> {prelims.get('keywords')}", style_normal))
        story.append(PageBreak())

    story.append(ResetPageNumberedCanvas())

    # Capítulo I
    c1 = data.get("capitulo1", {})
    story.append(Paragraph("CAPÍTULO I: INTRODUCCIÓN", style_header))
    story.append(Spacer(1, 1*cm))
    
    sections = [
        ("1.1 Realidad problemática", c1.get("realidad_problematica")),
        ("1.2 Antecedentes de la investigación", c1.get("antecedentes")),
        ("1.3 Justificación de la investigación", c1.get("justificacion")),
        ("1.4 Formulación del problema", c1.get("formulacion_problema")),
        ("1.5 Hipótesis", c1.get("hipotesis")),
        ("1.6 Objetivos", c1.get("objetivos")),
        ("1.7 Limitaciones del estudio", c1.get("limitaciones"))
    ]
    for title, text in sections:
        if text:
            story.append(Paragraph(title, style_section))
            story.append(Paragraph(get_clean_html_text(text), style_normal))
    story.append(PageBreak())

    # Capítulo II
    c2 = data.get("capitulo2", {})
    story.append(Paragraph("CAPÍTULO II: MARCO TEÓRICO", style_header))
    story.append(Spacer(1, 1*cm))
    if c2.get("bases_teoricas"):
        story.append(Paragraph("2.1 Bases teóricas", style_section))
        story.append(Paragraph(get_clean_html_text(c2.get("bases_teoricas")), style_normal))
    story.append(PageBreak())

    # Capítulo III
    c3 = data.get("capitulo3", {})
    story.append(Paragraph("CAPÍTULO III: MÉTODO", style_header))
    story.append(Spacer(1, 1*cm))
    
    sections_c3 = [
        ("3.1 Tipo de investigación", c3.get("tipo_investigacion")),
        ("3.2 Nivel de investigación", c3.get("nivel_investigacion")),
        ("3.3 Diseño de investigación", c3.get("diseno_investigacion"))
    ]
    for title, text in sections_c3:
        if text:
            story.append(Paragraph(title, style_section))
            story.append(Paragraph(text, style_normal))

    # Población, muestra y muestreo
    p_pm = c3.get("poblacion_muestra", {})
    if p_pm:
        story.append(Paragraph("3.4 Población, muestra y muestreo", style_section))
        for sub_t, key in [
            ("3.4.1 Población", "poblacion"),
            ("3.4.2 Muestra", "muestra"),
            ("3.4.3 Muestreo", "muestreo")
        ]:
            val = p_pm.get(key)
            if val:
                story.append(Paragraph(f"<b>{sub_t}</b>", style_normal))
                story.append(Paragraph(val, style_normal))

    # Variables y operacionalización
    vars_data = c3.get("variables", {})
    if vars_data:
        story.append(Paragraph("3.5 Variables", style_section))
        if vars_data.get("tipo"):
            story.append(Paragraph("<b>3.5.1 Tipo de Variables</b>", style_normal))
            story.append(Paragraph(vars_data.get("tipo"), style_normal))
            
        op_table_data = vars_data.get("operacionalizacion_tabla", [])
        if op_table_data:
            story.append(Paragraph("<b>3.5.2 Operacionalización de Variables</b>", style_normal))
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph("<i>Tabla 1. Matriz de Operacionalización de Variables</i>", style_normal))
            
            t_data = [['Variable', 'Def. Conceptual', 'Def. Operacional', 'Dimensiones', 'Indicadores', 'Escala']]
            for r in op_table_data:
                t_data.append([
                    r.get("variable", ""), r.get("definicion_conceptual", ""),
                    r.get("definicion_operacional", ""), r.get("dimensiones", ""),
                    r.get("indicadores", ""), r.get("escala_medicion", "")
                ])
            t = Table(t_data, colWidths=[2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm, 2.5*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.grey),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('GRID', (0,0), (-1,-1), 1, colors.black),
                ('FONTSIZE', (0,0), (-1,-1), 8),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.5*cm))

    # Técnicas e instrumentos
    ti_data = c3.get("tecnicas_instrumentos", {})
    if ti_data:
        story.append(Paragraph("3.6 Técnicas e instrumentos", style_section))
        if ti_data.get("descripcion"):
            story.append(Paragraph("<b>3.6.1 Técnicas e instrumentos de recolección</b>", style_normal))
            story.append(Paragraph(ti_data.get("descripcion"), style_normal))
        if ti_data.get("validacion_confiabilidad"):
            story.append(Paragraph("<b>3.6.2 Validación y confiabilidad</b>", style_normal))
            story.append(Paragraph(ti_data.get("validacion_confiabilidad"), style_normal))

    sections_c3_end = [
        ("3.7 Método de análisis de datos", c3.get("metodo_analisis")),
        ("3.8 Procedimiento realizado", c3.get("procedimiento")),
        ("3.9 Consideraciones éticas", c3.get("consideraciones_eticas"))
    ]
    for title, text in sections_c3_end:
        if text:
            story.append(Paragraph(title, style_section))
            story.append(Paragraph(text, style_normal))
    story.append(PageBreak())

    # Capítulo IV
    c4 = data.get("capitulo4", {})
    story.append(Paragraph("CAPÍTULO IV: RESULTADOS", style_header))
    story.append(Spacer(1, 1*cm))
    if c4.get("resultados"):
        story.append(Paragraph("4.1 Análisis descriptivo y contrastación de hipótesis", style_section))
        story.append(Paragraph(get_clean_html_text(c4.get("resultados")), style_normal))

    # Tables in Results
    if c4.get("resultados_tablas"):
        for idx, tab in enumerate(c4.get("resultados_tablas")):
            story.append(Paragraph(f"<b>{tab.get('titulo', f'Tabla {idx+2}')}</b>", style_normal))
            story.append(Spacer(1, 0.3*cm))
            cols = tab.get("columnas", [])
            filas = tab.get("filas", [])
            t_data = [cols]
            for r in filas:
                row_vals = [str(r.get(col, "")) for col in cols]
                t_data.append(row_vals)
            t = Table(t_data)
            t.setStyle(TableStyle([
                ('LINEABOVE', (0,0), (-1,0), 0.5, colors.black),
                ('LINEBELOW', (0,0), (-1,0), 0.5, colors.black),
                ('LINEBELOW', (0,-1), (-1,-1), 0.5, colors.black),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('FONTSIZE', (0,0), (-1,-1), 9),
            ]))
            story.append(t)
            story.append(Spacer(1, 0.5*cm))

    # Figures in Results
    if c4.get("resultados_figuras"):
        for fig in c4.get("resultados_figuras"):
            story.append(Spacer(1, 0.3*cm))
            story.append(Paragraph(f"<b>[ {fig.get('titulo', 'Figura')} ]</b>", style_normal))
            story.append(Paragraph(f"<i>{fig.get('titulo', 'Figura')}. {fig.get('descripcion', '')}</i>", style_normal))
            story.append(Spacer(1, 0.3*cm))
    story.append(PageBreak())

    # Capítulo V
    c5 = data.get("capitulo5", {})
    story.append(Paragraph("CAPÍTULO V: DISCUSIÓN Y RECOMENDACIONES", style_header))
    story.append(Spacer(1, 1*cm))
    if c5.get("discusion"):
        story.append(Paragraph("5.1 Discusión de los hallazgos", style_section))
        story.append(Paragraph(get_clean_html_text(c5.get("discusion")), style_normal))
    
    recs = c5.get("recomendaciones", [])
    if recs:
        story.append(Paragraph("5.2 Recomendaciones", style_section))
        for idx, rec in enumerate(recs):
            story.append(Paragraph(f"<b>{idx+1}.</b> {rec}", style_normal))
    story.append(PageBreak())

    # Conclusiones
    story.append(Paragraph("CONCLUSIONES", style_header))
    story.append(Spacer(1, 1*cm))
    concs = data.get("conclusiones", [])
    for idx, conc in enumerate(concs):
        story.append(Paragraph(f"<b>{idx+1}.</b> {conc}", style_normal))
    story.append(PageBreak())

    # Referencias
    story.append(Paragraph("REFERENCIAS BIBLIOGRÁFICAS", style_header))
    story.append(Spacer(1, 1*cm))
    refs = data.get("referencias", [])
    for ref in refs:
        story.append(Paragraph(ref, style_ref))
    story.append(PageBreak())

    # Anexos
    story.append(Paragraph("ANEXOS OBLIGATORIOS", style_header))
    story.append(Spacer(1, 1*cm))
    anexos = data.get("anexos", {})
    sections_an = [
        ("Anexo 1. Matriz de operacionalización de variables", anexos.get("anexo_1")),
        ("Anexo 2. Matriz de consistencia", anexos.get("anexo_2")),
        ("Anexo 3. Diagrama de Ishikawa", anexos.get("anexo_3")),
        ("Anexo 4. Árbol de problemas", anexos.get("anexo_4")),
        ("Anexo 5. Árbol de objetivos", anexos.get("anexo_5")),
        ("Anexo 6. Instrumentos de recolección de datos", anexos.get("anexo_6")),
        ("Anexo 7. Constancia de aplicación de instrumentos", anexos.get("anexo_7")),
        ("Anexo 8. Declaración de originalidad y conformidad", anexos.get("anexo_8"))
    ]
    for title, text in sections_an:
        if text:
            story.append(Paragraph(title, style_section))
            elements = parse_markdown_to_elements(text)
            append_markdown_elements_to_pdf_story(story, elements, style_normal, style_section)

    try:
        pdf.build(story, canvasmaker=NumberedCanvas)
    finally:
        for fpath in temp_files_to_clean:
            try: os.unlink(fpath)
            except: pass


def build_final_thesis_txt(data, output_path):
    meta = data.get("metadata", {})
    prelims = data.get("preliminares", {})
    c1 = data.get("capitulo1", {})
    c2 = data.get("capitulo2", {})
    c3 = data.get("capitulo3", {})
    c4 = data.get("capitulo4", {})
    c5 = data.get("capitulo5", {})
    refs = data.get("referencias", [])
    anexos = data.get("anexos", {})
    concs = data.get("conclusiones", [])
    
    with open(output_path, "w", encoding="utf-8") as f:
        template_styles = data.get("templateStyles", {}) or {}
        structure = template_styles.get("structure", {}) or {}
        institution = (structure.get("institution") or meta.get("universidad") or "UNIVERSIDAD NACIONAL DE TRUJILLO").upper()
        faculty = (structure.get("faculty") or meta.get("facultad") or "FACULTAD DE INGENIERÍA").upper()
        school = (structure.get("school") or meta.get("escuela") or "ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS").upper()

        f.write(f"{institution}\n{faculty}\n{school}\n\n")
        f.write(f"TEMA: {meta.get('titulo_proyecto', '').upper()}\n\n")
        f.write(f"Autor: {meta.get('nombre_autor', '')}\n")
        f.write(f"Asesor: Dr. {meta.get('nombre_asesor', '')}\n")
        f.write(f"Línea de Investigación: {meta.get('linea_investigacion', '')}\n")
        f.write(f"Ciudad: {meta.get('ciudad', '')} - Año: {meta.get('anio', '')}\n")
        f.write("=" * 60 + "\n\n")
        
        # Jurado
        f.write("JURADO DICTAMINADOR\n\n")
        f.write("Presidente: Dr. Roberto Carlos Medina\n")
        f.write("Secretario: Dr. Julio César Alvarez\n")
        f.write(f"Vocal (Asesor): Dr. {meta.get('nombre_asesor', '')}\n")
        f.write("=" * 60 + "\n\n")
        
        # Preliminares
        if prelims.get("dedicatoria"):
            f.write(f"DEDICATORIA\n\n{prelims.get('dedicatoria')}\n\n")
        if prelims.get("agradecimientos"):
            f.write(f"AGRADECIMIENTOS\n\n{prelims.get('agradecimientos')}\n\n")
        if prelims.get("presentacion"):
            f.write(f"PRESENTACIÓN\n\n{prelims.get('presentacion')}\n\n")
        if prelims.get("resumen"):
            f.write(f"RESUMEN\n\n{prelims.get('resumen')}\n")
            f.write(f"Palabras clave: {prelims.get('palabras_clave', '')}\n\n")
        if prelims.get("abstract"):
            f.write(f"ABSTRACT\n\n{prelims.get('abstract')}\n")
            f.write(f"Keywords: {prelims.get('keywords', '')}\n\n")
        f.write("=" * 60 + "\n\n")
            
        # Capitulo I
        f.write("CAPÍTULO I: INTRODUCCIÓN\n\n")
        sections = [
            ("1.1 Realidad problemática", c1.get("realidad_problematica")),
            ("1.2 Antecedentes de la investigación", c1.get("antecedentes")),
            ("1.3 Justificación de la investigación", c1.get("justificacion")),
            ("1.4 Formulación del problema", c1.get("formulacion_problema")),
            ("1.5 Hipótesis", c1.get("hipotesis")),
            ("1.6 Objetivos", c1.get("objetivos")),
            ("1.7 Limitaciones del estudio", c1.get("limitaciones"))
        ]
        for title, val in sections:
            f.write(f"{title}\n{get_clean_text(val)}\n\n")
        f.write("=" * 60 + "\n\n")

        # Capitulo II
        f.write("CAPÍTULO II: MARCO TEÓRICO\n\n")
        if c2.get("bases_teoricas"):
            f.write(f"2.1 Bases teóricas\n{c2.get('bases_teoricas')}\n\n")
        f.write("=" * 60 + "\n\n")

        # Capitulo III
        f.write("CAPÍTULO III: MÉTODO\n\n")
        f.write(f"3.1 Tipo de investigación\n{c3.get('tipo_investigacion','')}\n\n")
        f.write(f"3.2 Nivel de investigación\n{c3.get('nivel_investigacion','')}\n\n")
        f.write(f"3.3 Diseño de investigación\n{c3.get('diseno_investigacion','')}\n\n")
        
        p_pm = c3.get("poblacion_muestra", {})
        f.write(f"3.4 Población, muestra y muestreo\n")
        f.write(f"3.4.1 Población: {p_pm.get('poblacion','')}\n")
        f.write(f"3.4.2 Muestra: {p_pm.get('muestra','')}\n")
        f.write(f"3.4.3 Muestreo: {p_pm.get('muestreo','')}\n\n")
        
        vars_data = c3.get("variables", {})
        f.write(f"3.5 Variables\n")
        f.write(f"3.5.1 Tipo: {vars_data.get('tipo','')}\n\n")
        f.write("3.5.2 Operacionalización de Variables:\n")
        for row in vars_data.get("operacionalizacion_tabla", []):
            f.write(f"- Variable: {row.get('variable')}\n")
            f.write(f"  Def. Conceptual: {row.get('definicion_conceptual')}\n")
            f.write(f"  Def. Operacional: {row.get('definicion_operacional')}\n")
            f.write(f"  Dimensiones: {row.get('dimensiones')}\n")
            f.write(f"  Indicadores: {row.get('indicadores')}\n")
            f.write(f"  Escala: {row.get('escala_medicion')}\n")
        f.write("\n")
        
        ti_data = c3.get("tecnicas_instrumentos", {})
        f.write("3.6 Técnicas e instrumentos\n")
        f.write(f"3.6.1 Técnicas e instrumentos: {ti_data.get('descripcion','')}\n")
        f.write(f"3.6.2 Validación y confiabilidad: {ti_data.get('validacion_confiabilidad','')}\n\n")
        
        f.write(f"3.7 Método de análisis de datos\n{c3.get('metodo_analisis','')}\n\n")
        f.write(f"3.8 Procedimiento realizado\n{c3.get('procedimiento','')}\n\n")
        f.write(f"3.9 Consideraciones éticas\n{c3.get('consideraciones_eticas','')}\n\n")
        f.write("=" * 60 + "\n\n")

        # Capitulo IV
        f.write("CAPÍTULO IV: RESULTADOS\n\n")
        if c4.get("resultados"):
            f.write(f"4.1 Análisis descriptivo y contrastación de hipótesis\n{c4.get('resultados')}\n\n")
        if c4.get("resultados_tablas"):
            for idx, tab in enumerate(c4.get("resultados_tablas")):
                f.write(f"[{tab.get('titulo')}]\n")
                cols = tab.get("columnas", [])
                f.write(", ".join(cols) + "\n")
                for row in tab.get("filas", []):
                    f.write(", ".join([str(row.get(col, "")) for col in cols]) + "\n")
                f.write("\n")
        if c4.get("resultados_figuras"):
            for fig in c4.get("resultados_figuras"):
                f.write(f"[{fig.get('titulo')}]: {fig.get('descripcion')}\n\n")
        f.write("=" * 60 + "\n\n")

        # Capitulo V
        f.write("CAPÍTULO V: DISCUSIÓN Y RECOMENDACIONES\n\n")
        if c5.get("discusion"):
            f.write(f"5.1 Discusión de los hallazgos\n{c5.get('discusion')}\n\n")
        recs = c5.get("recomendaciones", [])
        if recs:
            f.write("5.2 Recomendaciones:\n")
            for idx, rec in enumerate(recs):
                f.write(f"{idx+1}. {rec}\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # Conclusiones
        f.write("CONCLUSIONES\n\n")
        for idx, conc in enumerate(concs):
            f.write(f"{idx+1}. {conc}\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # Referencias
        f.write("REFERENCIAS BIBLIOGRÁFICAS\n\n")
        for ref in refs:
            f.write(f"{ref}\n")
        f.write("\n" + "=" * 60 + "\n\n")

        # Anexos
        f.write("ANEXOS OBLIGATORIOS\n\n")
        for key, val in [
            ("Anexo 1. Matriz de operacionalización de variables", anexos.get("anexo_1")),
            ("Anexo 2. Matriz de consistencia", anexos.get("anexo_2")),
            ("Anexo 3. Diagrama de Ishikawa", anexos.get("anexo_3")),
            ("Anexo 4. Árbol de problemas", anexos.get("anexo_4")),
            ("Anexo 5. Árbol de objetivos", anexos.get("anexo_5")),
            ("Anexo 6. Instrumentos de recolección de datos", anexos.get("anexo_6")),
            ("Anexo 7. Constancia de aplicación de instrumentos", anexos.get("anexo_7")),
            ("Anexo 8. Declaración de originalidad y conformidad", anexos.get("anexo_8"))
        ]:
            f.write(f"{key}\n{val}\n\n")

# ---------------------------------------------------------------------------
# Article Export Functions (IMRyD format with Times New Roman & Strict Rules)
# ---------------------------------------------------------------------------

def build_article_docx(data, output_path):
    art = data.get("articulo", {})
    doc = Document()
    
    template_styles = data.get("templateStyles", {}) or {}
    
    # Margins for Article
    for section in doc.sections:
        apply_section_margins(section, template_styles, default_left=2.54)
        
        # Upper right corner page numbering in Header
        header = section.header
        header.is_linked_to_previous = False
        p_header = header.paragraphs[0]
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_page_number_to_footer(p_header)

    # --- Title ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(art.get("titulo", "ARTÍCULO CIENTÍFICO").upper())
    apply_text_formatting(run, font_name="Times New Roman", size=14, bold=True)
    set_article_paragraph_spacing(p, before=0, after=12, first_line_indent=0)

    # --- Authors ---
    p_auth = doc.add_paragraph()
    p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_auth = p_auth.add_run(art.get("autores", ""))
    apply_text_formatting(run_auth, font_name="Times New Roman", size=12, italic=True)
    set_article_paragraph_spacing(p_auth, before=0, after=18, first_line_indent=0)

    # --- Resumen / Abstract ---
    for lang_title, body_key, kw_key in [
        ("RESUMEN", "resumen", "palabras_clave"),
        ("ABSTRACT", "abstract", "keywords"),
    ]:
        body = art.get(body_key, "")
        kw   = art.get(kw_key, "")
        if body:
            p_t = doc.add_paragraph()
            run_t = p_t.add_run(lang_title)
            apply_text_formatting(run_t, font_name="Times New Roman", size=12, bold=True)
            set_article_paragraph_spacing(p_t, before=12, after=4, first_line_indent=0)

            p_b = doc.add_paragraph()
            run_b = p_b.add_run(body)
            apply_text_formatting(run_b, font_name="Times New Roman", size=11, italic=(lang_title == "ABSTRACT"))
            set_article_paragraph_spacing(p_b, before=0, after=6, first_line_indent=1.27)

            if kw:
                p_kw = doc.add_paragraph()
                kw_lbl = "Palabras clave: " if lang_title == 'RESUMEN' else "Keywords: "
                run_kw_lbl = p_kw.add_run(kw_lbl)
                apply_text_formatting(run_kw_lbl, font_name="Times New Roman", size=10, bold=True)
                run_kw_val = p_kw.add_run(kw)
                apply_text_formatting(run_kw_val, font_name="Times New Roman", size=10)
                set_article_paragraph_spacing(p_kw, before=0, after=12, first_line_indent=0)

    # --- Sections ---
    imryd_sections = [
        ("1. INTRODUCCIÓN",  art.get("introduccion") or art.get("introducción", "")),
        ("2. MÉTODOS",       art.get("metodos")      or art.get("métodos",      "")),
        ("3. RESULTADOS",    art.get("resultados",   "")),
        ("4. DISCUSIÓN",     art.get("discusion")    or art.get("discusión",    "")),
    ]
    for sec_title, sec_body in imryd_sections:
        if sec_body:
            p_sec = doc.add_paragraph()
            run_sec = p_sec.add_run(sec_title)
            apply_text_formatting(run_sec, font_name="Times New Roman", size=13, bold=True)
            set_article_paragraph_spacing(p_sec, before=18, after=6, first_line_indent=0)

            p_body = doc.add_paragraph()
            run_body = p_body.add_run(sec_body)
            apply_text_formatting(run_body, font_name="Times New Roman", size=12)
            set_article_paragraph_spacing(p_body, before=0, after=6, first_line_indent=1.27)

            # Draw structured tables if inside Results
            if "3. RESULTADOS" in sec_title and art.get("resultados_tablas"):
                for tab in art.get("resultados_tablas", []):
                    # Caption ABOVE the table (Times New Roman, 11 pt, bold)
                    p_cap = doc.add_paragraph()
                    run_cap = p_cap.add_run(tab.get("titulo", ""))
                    apply_text_formatting(run_cap, font_name="Times New Roman", size=11, bold=True)
                    set_article_paragraph_spacing(p_cap, before=12, after=4, first_line_indent=0)
                    
                    cols = tab.get("columnas", [])
                    filas = tab.get("filas", [])
                    table = doc.add_table(rows=1, cols=len(cols))
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Header row
                    hdr_cells = table.rows[0].cells
                    for idx, cname in enumerate(cols):
                        hdr_cells[idx].text = cname
                        apply_text_formatting(hdr_cells[idx].paragraphs[0].runs[0], font_name="Times New Roman", size=10, bold=True)
                        # Left align if text, otherwise right
                        hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Data rows
                    for r in filas:
                        row_cells = table.add_row().cells
                        for idx, cname in enumerate(cols):
                            val = r.get(cname, "")
                            row_cells[idx].text = str(val)
                            p_cell = row_cells[idx].paragraphs[0]
                            apply_text_formatting(p_cell.runs[0], font_name="Times New Roman", size=10)
                            # Align numbers to right, text to left
                            try:
                                float(str(val).replace(',','.'))
                                p_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            except ValueError:
                                p_cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                    apply_three_line_table_style(table)
                    doc.add_paragraph()  # spacer

            # Draw figures descriptions if inside Results
            if "3. RESULTADOS" in sec_title and art.get("resultados_figuras"):
                for fig in art.get("resultados_figuras", []):
                    p_fig = doc.add_paragraph()
                    p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_f = p_fig.add_run(f"[ {fig.get('titulo', 'Figura')} ]")
                    apply_text_formatting(run_f, font_name="Times New Roman", size=11, bold=True)
                    set_article_paragraph_spacing(p_fig, before=12, after=4, first_line_indent=0)
                    
                    # Caption BELOW in 10 pt italic
                    p_fig_cap = doc.add_paragraph()
                    p_fig_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_fc = p_fig_cap.add_run(f"{fig.get('titulo', 'Figura')}. {fig.get('descripcion', '')}")
                    apply_text_formatting(run_fc, font_name="Times New Roman", size=10, italic=True)
                    set_article_paragraph_spacing(p_fig_cap, before=4, after=12, first_line_indent=0)

    # --- Conclusiones ---
    conclusiones = art.get("conclusiones", [])
    if conclusiones:
        p_ct = doc.add_paragraph()
        run_ct = p_ct.add_run("5. CONCLUSIONES")
        apply_text_formatting(run_ct, font_name="Times New Roman", size=13, bold=True)
        set_article_paragraph_spacing(p_ct, before=18, after=6, first_line_indent=0)

        if isinstance(conclusiones, list):
            for idx, c in enumerate(conclusiones):
                p_c = doc.add_paragraph()
                run_n = p_c.add_run(f"{idx + 1}. ")
                apply_text_formatting(run_n, font_name="Times New Roman", bold=True)
                run_c = p_c.add_run(c)
                apply_text_formatting(run_c, font_name="Times New Roman", size=12)
                set_article_paragraph_spacing(p_c, before=0, after=4, first_line_indent=1.27)
        else:
            p_c = doc.add_paragraph()
            run_c = p_c.add_run(str(conclusiones))
            apply_text_formatting(run_c, font_name="Times New Roman", size=12)
            set_article_paragraph_spacing(p_c, before=0, after=4, first_line_indent=1.27)

    # --- Agradecimientos ---
    agr = art.get("agradecimientos", "")
    if agr:
        p_agr = doc.add_paragraph()
        run_agr = p_agr.add_run("AGRADECIMIENTOS")
        apply_text_formatting(run_agr, font_name="Times New Roman", size=13, bold=True)
        set_article_paragraph_spacing(p_agr, before=18, after=6, first_line_indent=0)
        
        p_b = doc.add_paragraph()
        run_b = p_b.add_run(agr)
        apply_text_formatting(run_b, font_name="Times New Roman", size=12)
        set_article_paragraph_spacing(p_b, before=0, after=6, first_line_indent=1.27)

    # --- Declaraciones Obligatorias ---
    declaraciones = art.get("declaraciones", {})
    if declaraciones:
        p_dec = doc.add_paragraph()
        run_dec = p_dec.add_run("DECLARACIONES OBLIGATORIAS")
        apply_text_formatting(run_dec, font_name="Times New Roman", size=13, bold=True)
        set_article_paragraph_spacing(p_dec, before=18, after=6, first_line_indent=0)
        
        for k, v in [
            ("Conflicto de intereses", declaraciones.get("conflicto_intereses", "No existe ningún tipo de conflicto de interés relacionado con la materia del trabajo")),
            ("Fuente de financiamiento", declaraciones.get("financiamiento", "Los autores no recibieron ningún patrocinio para llevar a cabo este estudio-artículo")),
            ("Contribución de autoría (Taxonomía CRediT)", declaraciones.get("contribucion_autores", "")),
            ("Disponibilidad de datos", declaraciones.get("disponibilidad_datos", "No aplica"))
        ]:
            if v:
                p_item = doc.add_paragraph()
                run_k = p_item.add_run(f"{k}: ")
                apply_text_formatting(run_k, font_name="Times New Roman", size=11, bold=True)
                run_v = p_item.add_run(str(v))
                apply_text_formatting(run_v, font_name="Times New Roman", size=11)
                set_article_paragraph_spacing(p_item, before=0, after=4, first_line_indent=0)

    # --- Referencias ---
    referencias = art.get("referencias", [])
    if referencias:
        p_rt = doc.add_paragraph()
        run_rt = p_rt.add_run("REFERENCIAS BIBLIOGRÁFICAS")
        apply_text_formatting(run_rt, font_name="Times New Roman", size=13, bold=True)
        set_article_paragraph_spacing(p_rt, before=18, after=8, first_line_indent=0)

        if isinstance(referencias, list):
            for ref in referencias:
                p_ref = doc.add_paragraph()
                p_ref.paragraph_format.left_indent = Cm(1.27)
                p_ref.paragraph_format.first_line_indent = Cm(-1.27)
                run_ref = p_ref.add_run(ref)
                apply_text_formatting(run_ref, font_name="Times New Roman", size=10)
                set_article_paragraph_spacing(p_ref, line_spacing=1.5, before=0, after=4)
        else:
            p_ref = doc.add_paragraph()
            run_ref = p_ref.add_run(str(referencias))
            apply_text_formatting(run_ref, font_name="Times New Roman", size=10)
            set_article_paragraph_spacing(p_ref, line_spacing=1.5, before=0, after=4)

    doc.save(output_path)


def build_article_pdf(data, output_path):
    art = data.get("articulo", {})
    template_styles = data.get("templateStyles", {}) or {}
    margins = template_styles.get("margins", {})

    pdf = SimpleDocTemplate(
        output_path,
        pagesize=letter,
        leftMargin=margins.get("left", 2.54)*cm,
        rightMargin=margins.get("right", 2.54)*cm,
        topMargin=margins.get("top", 2.54)*cm,
        bottomMargin=margins.get("bottom", 2.54)*cm
    )

    styles = getSampleStyleSheet()
    
    font_reg, font_bold, font_italic = get_reportlab_fonts(is_article=True)
    line_sp = get_line_spacing(is_article=True)
    
    style_normal  = ParagraphStyle('AN', parent=styles['Normal'], fontName=font_reg,    fontSize=12, leading=int(12 * line_sp * 1.2), alignment=4, spaceAfter=6, firstLineIndent=1.27*cm)
    style_title   = ParagraphStyle('AT', parent=styles['Normal'], fontName=font_bold,   fontSize=14, leading=20, alignment=1, spaceAfter=12)
    style_author  = ParagraphStyle('AA', parent=styles['Normal'], fontName=font_italic, fontSize=11, leading=15, alignment=1, spaceAfter=18)
    style_section = ParagraphStyle('AS', parent=styles['Normal'], fontName=font_bold,   fontSize=13, leading=18, alignment=0, spaceBefore=18, spaceAfter=6)
    style_caption = ParagraphStyle('AC', parent=styles['Normal'], fontName=font_bold,   fontSize=10, leading=14, alignment=1, spaceAfter=4)
    style_legend  = ParagraphStyle('AL', parent=styles['Normal'], fontName=font_italic, fontSize=9,  leading=12, alignment=1, spaceAfter=8)
    style_ref     = ParagraphStyle('AR', parent=styles['Normal'], fontName=font_reg,    fontSize=10, leading=14, alignment=4, leftIndent=1.27*cm, firstLineIndent=-1.27*cm, spaceAfter=4)

    story = []

    story.append(Paragraph(art.get("titulo", "ARTÍCULO CIENTÍFICO").upper(), style_title))
    story.append(Paragraph(art.get("autores", ""), style_author))
    story.append(Spacer(1, 0.5*cm))

    for lang_title, body_key, kw_key in [
        ("RESUMEN", "resumen", "palabras_clave"),
        ("ABSTRACT", "abstract", "keywords"),
    ]:
        body = art.get(body_key, "")
        kw   = art.get(kw_key, "")
        if body:
            story.append(Paragraph(lang_title, style_section))
            story.append(Paragraph(body, style_normal))
            if kw:
                kw_lbl = "Palabras clave" if lang_title == "RESUMEN" else "Keywords"
                story.append(Paragraph(f"<b>{kw_lbl}:</b> {kw}", style_normal))
            story.append(Spacer(1, 0.4*cm))

    imryd_sections = [
        ("1. INTRODUCCIÓN", art.get("introduccion") or art.get("introducción", "")),
        ("2. MÉTODOS",      art.get("metodos")      or art.get("métodos",      "")),
        ("3. RESULTADOS",   art.get("resultados",   "")),
        ("4. DISCUSIÓN",    art.get("discusion")    or art.get("discusión",    "")),
    ]
    for sec_title, sec_body in imryd_sections:
        if sec_body:
            story.append(Paragraph(sec_title, style_section))
            story.append(Paragraph(sec_body, style_normal))
            
            # Tables in Results (Three-line format)
            if "3. RESULTADOS" in sec_title and art.get("resultados_tablas"):
                for tab in art.get("resultados_tablas", []):
                    story.append(Paragraph(f"<b>{tab.get('titulo', '')}</b>", style_caption))
                    
                    cols = tab.get("columnas", [])
                    filas = tab.get("filas", [])
                    t_data = [cols]
                    for r in filas:
                        row_vals = []
                        for c in cols:
                            row_vals.append(str(r.get(c, "")))
                        t_data.append(row_vals)
                        
                    t = Table(t_data)
                    t.setStyle(TableStyle([
                        ('LINEABOVE', (0,0), (-1,0), 0.5, colors.black),
                        ('LINEBELOW', (0,0), (-1,0), 0.5, colors.black),
                        ('LINEBELOW', (0,-1), (-1,-1), 0.5, colors.black),
                        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                        ('FONTNAME', (0,0), (-1,-1), 'Times-Roman'),
                        ('FONTSIZE', (0,0), (-1,-1), 10),
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.5*cm))
                    
            # Figures in Results
            if "3. RESULTADOS" in sec_title and art.get("resultados_figuras"):
                for fig in art.get("resultados_figuras", []):
                    story.append(Spacer(1, 0.3*cm))
                    story.append(Paragraph(f"[ {fig.get('titulo', 'Figura')} ]", style_caption))
                    story.append(Paragraph(f"<i>{fig.get('titulo', 'Figura')}. {fig.get('descripcion', '')}</i>", style_legend))
                    story.append(Spacer(1, 0.3*cm))

    conclusiones = art.get("conclusiones", [])
    if conclusiones:
        story.append(Paragraph("5. CONCLUSIONES", style_section))
        if isinstance(conclusiones, list):
            for idx, c in enumerate(conclusiones):
                story.append(Paragraph(f"<b>{idx+1}.</b> {c}", style_normal))
        else:
            story.append(Paragraph(str(conclusiones), style_normal))
            
    agr = art.get("agradecimientos", "")
    if agr:
        story.append(Paragraph("AGRADECIMIENTOS", style_section))
        story.append(Paragraph(agr, style_normal))

    # --- Declaraciones Obligatorias ---
    declaraciones = art.get("declaraciones", {})
    if declaraciones:
        story.append(Paragraph("DECLARACIONES OBLIGATORIAS", style_section))
        for k, v in [
            ("Conflicto de intereses", declaraciones.get("conflicto_intereses", "No existe ningún tipo de conflicto de interés relacionado con la materia del trabajo")),
            ("Fuente de financiamiento", declaraciones.get("financiamiento", "Los autores no recibieron ningún patrocinio para llevar a cabo este estudio-artículo")),
            ("Contribución de autoría (Taxonomía CRediT)", declaraciones.get("contribucion_autores", "")),
            ("Disponibilidad de datos", declaraciones.get("disponibilidad_datos", "No aplica"))
        ]:
            if v:
                story.append(Paragraph(f"<b>{k}:</b> {v}", style_normal))

    referencias = art.get("referencias", [])
    if referencias:
        story.append(Paragraph("REFERENCIAS BIBLIOGRÁFICAS", style_section))
        if isinstance(referencias, list):
            for ref in referencias:
                story.append(Paragraph(ref, style_ref))
        else:
            story.append(Paragraph(str(referencias), style_ref))

    pdf.build(story, canvasmaker=ArticleNumberedCanvas)


def build_article_txt(data, output_path):
    art = data.get("articulo", {})

    with open(output_path, "w", encoding="utf-8") as f:
        sep = "=" * 70 + "\n\n"

        f.write(art.get("titulo", "ARTÍCULO CIENTÍFICO").upper() + "\n\n")
        f.write(f"Autores: {art.get('autores', '')}\n\n")
        f.write(sep)

        f.write("RESUMEN\n\n")
        f.write(f"{art.get('resumen', '')}\n")
        f.write(f"Palabras clave: {art.get('palabras_clave', '')}\n\n")

        f.write("ABSTRACT\n\n")
        f.write(f"{art.get('abstract', '')}\n")
        f.write(f"Keywords: {art.get('keywords', '')}\n\n")
        f.write(sep)

        for sec_title, body_key in [
            ("1. INTRODUCCIÓN", ("introduccion", "introducción")),
            ("2. MÉTODOS",      ("metodos",      "métodos")),
            ("3. RESULTADOS",   ("resultados",)),
            ("4. DISCUSIÓN",    ("discusion",    "discusión")),
        ]:
            body = ""
            for k in body_key:
                body = art.get(k, "")
                if body:
                    break
            if body:
                f.write(f"{sec_title}\n\n{body}\n\n")
                if "RESULTADOS" in sec_title:
                    if art.get("resultados_tablas"):
                        for tab in art.get("resultados_tablas", []):
                            f.write(f"[{tab.get('titulo')}]\n")
                            f.write(", ".join(tab.get("columnas", [])) + "\n")
                            for row in tab.get("filas", []):
                                f.write(", ".join([str(row.get(col, "")) for col in tab.get("columnas", [])]) + "\n")
                            f.write("\n")
                    if art.get("resultados_figuras"):
                        for fig in art.get("resultados_figuras", []):
                            f.write(f"[{fig.get('titulo')}]: {fig.get('descripcion')}\n\n")
                f.write(sep)

        conclusiones = art.get("conclusiones", [])
        if conclusiones:
            f.write("5. CONCLUSIONES\n\n")
            if isinstance(conclusiones, list):
                for idx, c in enumerate(conclusiones):
                    f.write(f"{idx+1}. {c}\n")
            else:
                f.write(str(conclusiones) + "\n")
            f.write("\n" + sep)
            
        agr = art.get("agradecimientos", "")
        if agr:
            f.write("AGRADECIMIENTOS\n\n" + agr + "\n\n" + sep)

        # --- Declaraciones Obligatorias ---
        declaraciones = art.get("declaraciones", {})
        if declaraciones:
            f.write("DECLARACIONES OBLIGATORIAS\n\n")
            for k, v in [
                ("Conflicto de intereses", declaraciones.get("conflicto_intereses", "No existe ningún tipo de conflicto de interés relacionado con la materia del trabajo")),
                ("Fuente de financiamiento", declaraciones.get("financiamiento", "Los autores no recibieron ningún patrocinio para llevar a cabo este estudio-artículo")),
                ("Contribución de autoría (Taxonomía CRediT)", declaraciones.get("contribucion_autores", "")),
                ("Disponibilidad de datos", declaraciones.get("disponibilidad_datos", "No aplica"))
            ]:
                if v:
                    f.write(f"{k}: {v}\n")
            f.write("\n" + sep)

        referencias = art.get("referencias", [])
        if referencias:
            f.write("REFERENCIAS BIBLIOGRÁFICAS\n\n")
            if isinstance(referencias, list):
                for ref in referencias:
                    f.write(f"{ref}\n")
            else:
                f.write(str(referencias) + "\n")


# ---------------------------------------------------------------------------
# Main Runner
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Uso: python generate_docs.py <json_path> <output_path> <format>")
        sys.exit(1)
        
    json_path = sys.argv[1]
    output_path = sys.argv[2]
    out_format = sys.argv[3].lower()
    
    # Load JSON data
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Load template overrides if available
    template_styles = data.get("templateStyles", {})
    if template_styles:
        GLOBAL_FONT_FAMILY = template_styles.get("fontFamily")
        GLOBAL_LINE_SPACING = template_styles.get("lineSpacing")
    
    # Detect document type: article vs thesis vs final thesis
    is_article = "articulo" in data
    is_final_thesis = data.get("is_final_thesis", False)

    if out_format == "docx":
        if is_article:
            build_article_docx(data, output_path)
        elif is_final_thesis:
            build_final_thesis_docx(data, output_path)
        else:
            build_docx(data, output_path)
    elif out_format == "pdf":
        if is_article:
            build_article_pdf(data, output_path)
        elif is_final_thesis:
            build_final_thesis_pdf(data, output_path)
        else:
            build_pdf(data, output_path)
    elif out_format == "txt":
        if is_article:
            build_article_txt(data, output_path)
        elif is_final_thesis:
            build_final_thesis_txt(data, output_path)
        else:
            build_txt(data, output_path)
    else:
        print(f"Formato no soportado: {out_format}")
        sys.exit(1)
        
    print(f"Archivo generado exitosamente en {output_path}")
