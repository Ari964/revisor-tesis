import sys
import json
import os
import zipfile
import base64
from docx import Document

def extract_first_image(docx_path):
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            file_list = z.namelist()
            media_files = [f for f in file_list if f.startswith('word/media/')]
            image_extensions = ('.png', '.jpeg', '.jpg', '.gif', '.bmp')
            image_files = [f for f in media_files if f.lower().endswith(image_extensions)]
            if image_files:
                # Sort to get the first image (usually image1.ext)
                first_img = sorted(image_files, key=lambda x: x.lower())[0]
                data = z.read(first_img)
                encoded = base64.b64encode(data).decode('utf-8')
                ext = first_img.split('.')[-1]
                return {
                    "base64": encoded,
                    "ext": ext
                }
    except Exception as e:
        pass
    return None

def parse_docx(file_path):
    if not os.path.exists(file_path):
        return {"error": f"File not found: {file_path}"}

    try:
        doc = Document(file_path)
    except Exception as e:
        return {"error": f"Failed to load docx: {str(e)}"}

    # 1. Margins (default first section)
    margins = {"top": 2.54, "bottom": 2.54, "left": 2.54, "right": 2.54}
    if doc.sections:
        section = doc.sections[0]
        try:
            margins["top"] = round(section.top_margin.cm, 2)
        except: pass
        try:
            margins["bottom"] = round(section.bottom_margin.cm, 2)
        except: pass
        try:
            margins["left"] = round(section.left_margin.cm, 2)
        except: pass
        try:
            margins["right"] = round(section.right_margin.cm, 2)
        except: pass

    # 2. Font and Line Spacing analysis
    font_names = {}
    line_spacings = []
    headings = []
    full_text = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        full_text.append(txt)

        # Headings detection
        if p.style and p.style.name.startswith("Heading"):
            headings.append(txt)
        elif txt.isupper() and (txt.startswith("CAPÍTULO") or txt.startswith("INTRODUCCIÓN") or txt.startswith("MÉTODO") or txt.startswith("REFERENCIAS") or txt.startswith("ANEXOS")):
            headings.append(txt)

        # Line spacing
        if p.paragraph_format.line_spacing is not None:
            # line_spacing can be float (e.g. 1.5, 2.0) or Pt (absolute spacing)
            ls = p.paragraph_format.line_spacing
            if isinstance(ls, float):
                line_spacings.append(ls)
            elif hasattr(ls, 'pt'):
                # Spacing in points, estimate ratio based on font size (usually 12pt)
                line_spacings.append(round(ls.pt / 12.0, 1))

        # Fonts in runs
        for run in p.runs:
            if run.font.name:
                font_names[run.font.name] = font_names.get(run.font.name, 0) + 1

    # Most common font
    font_family = "Arial"
    if font_names:
        font_family = max(font_names, key=font_names.get)

    # Average or most common line spacing
    line_spacing = 1.5
    if line_spacings:
        line_spacing = round(sum(line_spacings) / len(line_spacings), 2)
        # Snap to common spacings
        if abs(line_spacing - 1.0) < 0.15: line_spacing = 1.0
        elif abs(line_spacing - 1.15) < 0.1: line_spacing = 1.15
        elif abs(line_spacing - 1.5) < 0.2: line_spacing = 1.5
        elif abs(line_spacing - 2.0) < 0.2: line_spacing = 2.0

    # Table styles
    table_styles = {}
    for table in doc.tables:
        if table.style and table.style.name:
            table_styles[table.style.name] = table_styles.get(table.style.name, 0) + 1
    
    table_style = "Normal Table"
    if table_styles:
        table_style = max(table_styles, key=table_styles.get)

    logo_data = extract_first_image(file_path)

    return {
        "text": "\n".join(full_text[:500]), # Limit extracted text preview to avoid payload issues
        "full_text": "\n".join(full_text),
        "styles": {
            "fontFamily": font_family,
            "lineSpacing": line_spacing,
            "margins": margins,
            "tableStyle": table_style,
            "headings": headings[:50] # Keep a reasonable list of outline items
        },
        "logo": logo_data
    }

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(json.dumps({"error": "No file path provided"}))
        sys.exit(1)

    file_path = sys.argv[1]
    result = parse_docx(file_path)
    print(json.dumps(result, ensure_ascii=False, indent=2))
